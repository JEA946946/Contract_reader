from __future__ import annotations

import json
import logging
from pathlib import Path
from sqlalchemy.orm import Session

from app.parsers.base import ParsedPriceRow, ParsedMenuRow, ParsedTransportRow
from app.parsers.excel_parser import ExcelParser, CsvParser
from app.parsers.pdf_parser import PdfParser
from app.parsers.word_parser import WordParser
from app.parsers.email_parser import EmailParser
from app.parsers.text_parser import TextParser
from app.parsers.ai_parser import AiParser
from app.parsers.restaurant_ai_parser import RestaurantAiParser
from app.parsers.transportation_ai_parser import TransportationAiParser
from app.parsers.category_detector import detect_category_from_text, detect_category_from_pdf
from app.config import settings
from app.services.normalizer import save_parsed_rows
from app.utils import detect_file_type
from app.models import Document

logger = logging.getLogger(__name__)


PARSERS_BY_TYPE = {
    "xlsx": ExcelParser(),
    "xls": ExcelParser(),
    "csv": CsvParser(),
    "pdf": PdfParser(),
    "docx": WordParser(),
    "doc": WordParser(),
    "eml": EmailParser(),
    "msg": EmailParser(),
    "txt": TextParser(),
    "text": TextParser(),
}


def _extract_email_text(file_path: Path, file_type: str) -> str | None:
    """Extract subject + body text from an email for category detection."""
    try:
        if file_type == "eml":
            import email as email_mod
            from email import policy as email_policy
            import re as _re

            with open(file_path, "rb") as f:
                msg = email_mod.message_from_binary_file(f, policy=email_policy.default)

            parts = []
            subject = msg.get("Subject", "")
            if subject:
                parts.append(subject)

            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content:
                    # Strip HTML tags if HTML body
                    if body.get_content_type() == "text/html":
                        content = _re.sub(r"<[^>]+>", " ", content)
                    parts.append(content)

            return "\n".join(parts) if parts else None

        elif file_type == "msg":
            import extract_msg

            msg = extract_msg.Message(str(file_path))
            parts = []
            if msg.subject:
                parts.append(msg.subject)
            if msg.body:
                parts.append(msg.body)
            return "\n".join(parts) if parts else None
    except Exception as e:
        logger.warning("Email text extraction failed: %s", e)
    return None


def _extract_attachment_from_email(file_path: Path, file_type: str) -> Path | None:
    """Extract the first meaningful attachment (PDF/XLSX/DOCX) from an email.

    Returns path to a temp file. Caller is responsible for cleanup.
    """
    attachments = _extract_all_attachments_from_email(file_path, file_type)
    return attachments[0] if attachments else None


def _extract_all_attachments_from_email(file_path: Path, file_type: str) -> list[Path]:
    """Extract all meaningful attachments (PDF/XLSX/DOCX) from an email.

    Returns list of paths to temp files. Caller is responsible for cleanup.
    """
    ATTACHMENT_TYPES = {".pdf", ".xlsx", ".xls", ".docx", ".doc"}
    results: list[Path] = []

    try:
        if file_type == "eml":
            import email as email_mod
            from email import policy as email_policy

            with open(file_path, "rb") as f:
                msg = email_mod.message_from_binary_file(f, policy=email_policy.default)

            att_idx = 0
            for part in msg.walk():
                if part.get_content_disposition() in ("attachment", "inline") and part.get_filename():
                    filename = part.get_filename()
                    if not filename:
                        continue
                    suffix = Path(filename).suffix.lower()
                    if suffix not in ATTACHMENT_TYPES:
                        continue
                    data = part.get_payload(decode=True)
                    if not data:
                        continue
                    temp_file = file_path.parent / f"{file_path.stem}_extracted_{att_idx}{suffix}"
                    temp_file.write_bytes(data)
                    results.append(temp_file)
                    att_idx += 1

        elif file_type == "msg":
            import extract_msg

            msg = extract_msg.Message(str(file_path))
            att_idx = 0
            for attachment in msg.attachments:
                if not attachment.longFilename:
                    continue
                suffix = Path(attachment.longFilename).suffix.lower()
                if suffix not in ATTACHMENT_TYPES:
                    continue
                temp_file = file_path.parent / f"{file_path.stem}_extracted_{att_idx}{suffix}"
                temp_file.write_bytes(attachment.data)
                results.append(temp_file)
                att_idx += 1

    except Exception as e:
        logger.warning("Email attachment extraction failed: %s", e)
    return results


def _detect_category(file_path: Path, file_type: str) -> str:
    """Detect whether a document is hotel or restaurant."""
    try:
        if file_type == "pdf":
            return detect_category_from_pdf(file_path)
        elif file_type in ("txt", "csv", "text"):
            text = file_path.read_text(encoding="utf-8", errors="replace")
            return detect_category_from_text(text)
        elif file_type in ("docx", "doc"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            return detect_category_from_text(text)
        elif file_type in ("xlsx", "xls"):
            import pandas as pd
            dfs = pd.read_excel(file_path, sheet_name=None, header=None)
            text_parts = []
            for sheet_name, df in dfs.items():
                for _, row in df.iterrows():
                    cells = [str(c) if pd.notna(c) else "" for c in row]
                    if any(cells):
                        text_parts.append(" ".join(cells))
            return detect_category_from_text("\n".join(text_parts))
        elif file_type in ("eml", "msg"):
            text = _extract_email_text(file_path, file_type)
            if text:
                return detect_category_from_text(text)
    except Exception as e:
        logger.warning("Category detection failed: %s", e)
    return "hotel"


def _extract_text_for_ai(file_path: Path, file_type: str) -> str | None:
    """Extract text from a document for AI parsing."""
    if file_type in ("docx", "doc"):
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    text_parts.append(" | ".join(cells))
        return "\n".join(text_parts)
    elif file_type in ("txt", "csv", "text"):
        return file_path.read_text(encoding="utf-8", errors="replace")
    elif file_type in ("xlsx", "xls"):
        import pandas as pd
        dfs = pd.read_excel(file_path, sheet_name=None, header=None)
        text_parts = []
        for sheet_name, df in dfs.items():
            text_parts.append(f"--- Sheet: {sheet_name} ---")
            for _, row in df.iterrows():
                cells = [str(c) if pd.notna(c) else "" for c in row]
                if any(cells):
                    text_parts.append(" | ".join(cells))
        return "\n".join(text_parts)
    return None


def parse_restaurant_document(file_path: Path, document: Document, db: Session) -> list[ParsedMenuRow]:
    """Parse a restaurant document and return menu rows."""
    file_type = detect_file_type(file_path.name)
    rows: list[ParsedMenuRow] = []

    try:
        parser = RestaurantAiParser()
        if file_type == "pdf":
            rows = parser.parse_pdf(file_path)
        elif file_type in ("eml", "msg"):
            # Extract all attachments from email (Phase 4.2)
            att_paths = _extract_all_attachments_from_email(file_path, file_type)
            for att_path in att_paths:
                try:
                    att_type = detect_file_type(att_path.name)
                    if att_type == "pdf":
                        rows.extend(parser.parse_pdf(att_path))
                    else:
                        text = _extract_text_for_ai(att_path, att_type)
                        if text:
                            rows.extend(parser.parse_text(text))
                finally:
                    att_path.unlink(missing_ok=True)
            # Fallback: try email body text
            if not rows:
                email_text = _extract_email_text(file_path, file_type)
                if email_text:
                    rows = parser.parse_text(email_text)
        else:
            text = _extract_text_for_ai(file_path, file_type)
            if text:
                rows = parser.parse_text(text)
    except Exception as e:
        document.notes = (document.notes or "") + f" | Restaurant parser error: {str(e)}"

    if not rows:
        document.notes = (document.notes or "") + " | No menu data found"

    rows = deduplicate_menu_rows(rows)
    document.status = "pending_review"
    document.row_count = len(rows)
    return rows


def parse_transport_document(file_path: Path, document: Document, db: Session) -> list[ParsedTransportRow]:
    """Parse a transportation document and return transport rows."""
    file_type = detect_file_type(file_path.name)
    rows: list[ParsedTransportRow] = []

    try:
        parser = TransportationAiParser()
        if file_type == "pdf":
            rows = parser.parse_pdf(file_path)
        elif file_type in ("eml", "msg"):
            # Extract all attachments from email (Phase 4.2)
            att_paths = _extract_all_attachments_from_email(file_path, file_type)
            for att_path in att_paths:
                try:
                    att_type = detect_file_type(att_path.name)
                    if att_type == "pdf":
                        rows.extend(parser.parse_pdf(att_path))
                    else:
                        text = _extract_text_for_ai(att_path, att_type)
                        if text:
                            rows.extend(parser.parse_text(text))
                finally:
                    att_path.unlink(missing_ok=True)
            if not rows:
                email_text = _extract_email_text(file_path, file_type)
                if email_text:
                    rows = parser.parse_text(email_text)
        else:
            text = _extract_text_for_ai(file_path, file_type)
            if text:
                rows = parser.parse_text(text)
    except Exception as e:
        document.notes = (document.notes or "") + f" | Transport parser error: {str(e)}"

    if not rows:
        document.notes = (document.notes or "") + " | No transport data found"

    rows = deduplicate_transport_rows(rows)
    document.status = "pending_review"
    document.row_count = len(rows)
    return rows


def _run_ai_parse(file_path: Path, file_type: str, document: Document) -> list[ParsedPriceRow]:
    """Run full AI parsing (no rule-based context)."""
    rows: list[ParsedPriceRow] = []
    try:
        ai_parser = AiParser()
        if file_type in ("pdf",):
            rows = ai_parser.parse_pdf(file_path)
        elif file_type in ("docx", "doc"):
            text = _extract_text_for_ai(file_path, file_type)
            if text:
                rows = ai_parser.parse_text(text)
        elif file_type in ("txt", "csv", "text"):
            rows = ai_parser.parse(file_path)
        elif file_type in ("eml", "msg"):
            rows = EmailParser().parse(file_path)
        else:
            document.notes = (document.notes or "") + " | AI fallback: unsupported format"
    except Exception as e:
        document.notes = (document.notes or "") + f" | AI parser error: {str(e)}"
    return rows


def _run_ai_with_context(
    file_path: Path,
    file_type: str,
    rule_rows: list[ParsedPriceRow],
    document: Document,
) -> list[ParsedPriceRow]:
    """Run AI parsing seeded with partial rule-based context (hybrid mode).

    Provides the rule-based extraction as additional context so the AI can
    fill gaps rather than starting from scratch.
    """
    import json as _json

    ai_rows = _run_ai_parse(file_path, file_type, document)

    if not ai_rows:
        # AI returned nothing — use what rule-based found
        return rule_rows

    # Merge: prefer AI rows but keep rule-based rows that AI missed
    # (e.g. if AI missed a hotel that rule-based caught)
    ai_keys = set()
    for r in ai_rows:
        ai_keys.add((
            (r.accommodation or "").strip().upper(),
            (r.room_desc or "").strip().upper(),
            (r.meal_plan or "").strip().upper(),
            (r.season_code or "").strip().upper(),
        ))

    merged = list(ai_rows)
    for r in rule_rows:
        key = (
            (r.accommodation or "").strip().upper(),
            (r.room_desc or "").strip().upper(),
            (r.meal_plan or "").strip().upper(),
            (r.season_code or "").strip().upper(),
        )
        if key not in ai_keys:
            merged.append(r)

    logger.info("Hybrid merge: %d AI rows + %d unique rule rows = %d total",
                len(ai_rows), len(merged) - len(ai_rows), len(merged))
    return merged


def parse_document(file_path: Path, document: Document, db: Session) -> list[ParsedPriceRow]:
    """Parse a document and return rows without saving to DB.

    Sets document.status to 'pending_review'.
    """
    file_type = detect_file_type(file_path.name)

    # Detect category if not already set
    if not document.document_category or document.document_category == "hotel":
        detected = _detect_category(file_path, file_type)
        document.document_category = detected
        logger.info("Document %d category detected: %s", document.id, detected)

    # If restaurant, delegate to restaurant parser
    if document.document_category == "restaurant":
        menu_rows = parse_restaurant_document(file_path, document, db)
        # Store as parsed_rows_json using menu format
        document.parsed_rows_json = rows_to_json(menu_rows)
        return []  # Return empty hotel rows

    # If transportation, delegate to transport parser
    if document.document_category == "transportation":
        transport_rows = parse_transport_document(file_path, document, db)
        document.parsed_rows_json = rows_to_json(transport_rows)
        return []  # Return empty hotel rows

    # Optimized extraction pipeline: truncation → cache → rule-based → Haiku → Sonnet
    from app.parsers.optimized_processor import OptimizedDocumentProcessor

    processor = OptimizedDocumentProcessor(
        anthropic_api_key=settings.anthropic_api_key,
        redis_url=settings.redis_url,
        max_pages=settings.max_pages_per_document,
        max_chars=settings.max_chars_per_ai_call,
        enable_sonnet_fallback=True,
        validation_threshold=settings.validation_threshold,
        cache_ttl_days=settings.redis_cache_ttl_days,
    )

    result = processor.process_document(str(file_path), document.content_hash)
    rows = result.get("rows", [])
    method = result.get("extraction_method", "unknown")
    confidence = result.get("confidence", 0.0)
    logger.info(
        "Document %d: optimized pipeline — method=%s, confidence=%.2f, rows=%d",
        document.id, method, confidence, len(rows),
    )
    if method != "rule_based":
        document.notes = (document.notes or "") + f" | extraction_method={method}, confidence={confidence:.2f}"

    if not rows:
        document.notes = (document.notes or "") + " | No price data found"

    rows = deduplicate_rows(rows)
    rows = collapse_hotel_rows(rows)

    document.status = "pending_review"
    document.row_count = len(rows)
    return rows


def _row_dedup_key(row: ParsedPriceRow) -> tuple:
    """Build a hashable key from all meaningful fields of a row."""
    date_ranges_key = tuple(
        (dr.date_from, dr.date_to) for dr in row.date_ranges
    )
    return (
        (row.accommodation or "").strip().upper(),
        (row.city or "").strip().upper(),
        (row.room_desc or "").strip().upper(),
        row.double_price,
        row.single_price,
        row.twin_price,
        row.triple_price,
        row.quadruple_price,
        (row.meal_plan or "").strip().upper(),
        (row.season_code or "").strip().upper(),
        row.stars,
        (row.hotel_type or "").strip().upper(),
        (row.fit_git or "").strip().upper(),
        date_ranges_key,
    )


def deduplicate_rows(rows: list[ParsedPriceRow]) -> list[ParsedPriceRow]:
    """Remove exact duplicate rows, preserving order."""
    seen: set = set()
    unique: list[ParsedPriceRow] = []
    for row in rows:
        key = _row_dedup_key(row)
        if key not in seen:
            seen.add(key)
            unique.append(row)
    if len(unique) < len(rows):
        logger.info(f"Deduplication removed {len(rows) - len(unique)} duplicate rows ({len(rows)} → {len(unique)})")
    return unique


def collapse_hotel_rows(rows: list[ParsedPriceRow]) -> list[ParsedPriceRow]:
    """Merge rows with the same hotel/room/meal/season that have complementary prices.

    E.g., two rows for "Superior BB Annual" — one with dbl+sgl+twn, another with trp
    — get collapsed into a single row with all four prices.
    Only merges when prices don't conflict (both rows have different non-null values
    for the same column → keep separate).
    """
    from dataclasses import replace

    def collapse_key(row: ParsedPriceRow) -> tuple:
        date_ranges_key = tuple(
            (dr.date_from, dr.date_to) for dr in row.date_ranges
        )
        return (
            (row.accommodation or "").strip().upper(),
            (row.city or "").strip().upper(),
            (row.room_desc or "").strip().upper(),
            (row.meal_plan or "").strip().upper(),
            (row.season_code or "").strip().upper(),
            row.stars,
            (row.hotel_type or "").strip().upper(),
            (row.fit_git or "").strip().upper(),
            date_ranges_key,
        )

    # Group rows by collapse key, preserving order
    groups: dict[tuple, list[ParsedPriceRow]] = {}
    order: list[tuple] = []
    for row in rows:
        key = collapse_key(row)
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(row)

    def _clean_misplaced_prices(group: list[ParsedPriceRow]) -> list[ParsedPriceRow]:
        """Pre-clean a group of rows that look like occupancy variants without proper labels.

        Detects patterns like:
          Row A: dbl=1050, sgl=1050           (double occupancy row)
          Row B: dbl=1100, sgl=1100, twn=1100 (twin occupancy - dbl/sgl are misplaced)
          Row C: sgl=1650, trp=1650           (triple occupancy - sgl is misplaced)

        Strips the misplaced prices so they don't conflict during merge.
        """
        if len(group) < 2:
            return group

        cleaned = []
        for row in group:
            updates = {}
            # If twin_price == double_price and they're non-null, this is a "twin row"
            # — the double_price is actually the twin rate misplaced
            if (row.twin_price is not None and row.double_price is not None
                    and row.twin_price == row.double_price):
                updates["double_price"] = None
                if row.single_price == row.twin_price:
                    updates["single_price"] = None

            # If triple_price == single_price and they're non-null, this is a "triple row"
            if (row.triple_price is not None and row.single_price is not None
                    and row.triple_price == row.single_price):
                updates["single_price"] = None
                if row.double_price is not None and row.double_price == row.triple_price:
                    updates["double_price"] = None
                if row.twin_price is not None and row.twin_price == row.triple_price:
                    updates["twin_price"] = None

            # If quadruple_price matches other prices, it's a "quad row"
            if (row.quadruple_price is not None and row.double_price is not None
                    and row.quadruple_price == row.double_price):
                updates["double_price"] = None
                if row.single_price == row.quadruple_price:
                    updates["single_price"] = None
                if row.twin_price == row.quadruple_price:
                    updates["twin_price"] = None
                if row.triple_price == row.quadruple_price:
                    updates["triple_price"] = None

            if updates:
                cleaned.append(replace(row, **updates))
            else:
                cleaned.append(row)
        return cleaned

    result: list[ParsedPriceRow] = []
    total_collapsed = 0
    for key in order:
        group = groups[key]
        if len(group) == 1:
            result.append(group[0])
            continue

        # Pre-clean misplaced prices (e.g. twin rate in double_price column)
        group = _clean_misplaced_prices(group)

        # Try to merge all rows in the group into one
        merged = group[0]
        conflict = False
        for other in group[1:]:
            # Check for conflicts (both have different non-null values for same price)
            for attr in ("double_price", "single_price", "twin_price", "triple_price", "quadruple_price"):
                val_m = getattr(merged, attr)
                val_o = getattr(other, attr)
                if val_m is not None and val_o is not None and val_m != val_o:
                    conflict = True
                    break
            if conflict:
                break

            # Merge: fill in missing prices from the other row
            updates = {}
            for attr in ("double_price", "single_price", "twin_price", "triple_price", "quadruple_price"):
                if getattr(merged, attr) is None and getattr(other, attr) is not None:
                    updates[attr] = getattr(other, attr)
            # Also merge notes
            other_note = (other.note or "").strip()
            if other_note and other_note not in (merged.note or ""):
                current = (merged.note or "").strip()
                updates["note"] = f"{current}; {other_note}" if current else other_note
            if updates:
                merged = replace(merged, **updates)

        if conflict:
            # Can't merge — keep all rows as-is
            result.extend(group)
        else:
            result.append(merged)
            total_collapsed += len(group) - 1

    if total_collapsed:
        logger.info(f"Row collapse merged {total_collapsed} rows ({len(rows)} → {len(result)})")
    return result


def _has_meaningful_data(rows: list[ParsedPriceRow]) -> bool:
    """Check if any row has at least one price value."""
    return any(r.double_price or r.single_price or r.twin_price or r.triple_price or r.quadruple_price for r in rows)


def _extraction_quality_score(rows: list[ParsedPriceRow]) -> float:
    """Score the quality of rule-based extraction from 0.0 to 1.0.

    Weights:
    - 50%: proportion of rows with at least one price
    - 30%: proportion of rows with a hotel/accommodation name
    - 20%: proportion of rows with at least one date range

    Returns 0.0 for empty rows.
    """
    if not rows:
        return 0.0

    n = len(rows)
    has_price = sum(
        1 for r in rows
        if any([r.double_price, r.single_price, r.twin_price, r.triple_price, r.quadruple_price])
    )
    has_name = sum(1 for r in rows if (r.accommodation or "").strip())
    has_dates = sum(1 for r in rows if r.date_ranges)

    score = (
        0.5 * (has_price / n)
        + 0.3 * (has_name / n)
        + 0.2 * (has_dates / n)
    )
    return round(score, 3)


def process_document(file_path: Path, document: Document, db: Session) -> int:
    """Parse and immediately save — backward-compatible wrapper."""
    rows = parse_document(file_path, document, db)
    if rows:
        count = save_parsed_rows(rows, document, db)
        document.row_count = count
        document.status = "completed"
        return count
    else:
        document.status = "completed"
        document.row_count = 0
        return 0


def _row_to_dict(row: ParsedPriceRow) -> dict:
    """Convert a ParsedPriceRow to a serializable dict matching ParsedRowSchema shape."""
    return {
        "accommodation": row.accommodation or "",
        "city": row.city or "",
        "room_desc": row.room_desc,
        "double_price": float(row.double_price) if row.double_price is not None else None,
        "single_price": float(row.single_price) if row.single_price is not None else None,
        "twin_price": float(row.twin_price) if row.twin_price is not None else None,
        "triple_price": float(row.triple_price) if row.triple_price is not None else None,
        "quadruple_price": float(row.quadruple_price) if row.quadruple_price is not None else None,
        "stars": row.stars,
        "hotel_type": row.hotel_type,
        "meal_plan": row.meal_plan,
        "fit_git": row.fit_git,
        "season_code": row.season_code,
        "baby_discount": row.baby_discount,
        "child_discount": row.child_discount,
        "date_ranges": [
            {
                "date_from": dr.date_from.isoformat() if dr.date_from else None,
                "date_to": dr.date_to.isoformat() if dr.date_to else None,
            }
            for dr in row.date_ranges
        ],
        "min_stay": row.min_stay,
        "note": row.note,
        "address": row.address,
        "phone": row.phone,
        "email": row.email,
    }


def _transport_row_to_dict(row: ParsedTransportRow) -> dict:
    """Convert a ParsedTransportRow to a serializable dict."""
    return {
        "code": row.code or "",
        "price": float(row.price) if row.price is not None else None,
        "company_name": row.company_name or "",
        "company_code": row.company_code or "",
        "product": row.product or "",
        "bus_size": row.bus_size,
        "service_type": row.service_type,
        "days": row.days,
        "route_description": row.route_description or "",
        "note": row.note or "",
        "city": row.city or "",
    }


def rows_to_json(rows: list) -> str:
    """Serialize parsed rows to JSON for caching. Handles hotel, restaurant, and transport rows."""
    if rows and isinstance(rows[0], ParsedMenuRow):
        return json.dumps([_menu_row_to_dict(r) for r in rows])
    if rows and isinstance(rows[0], ParsedTransportRow):
        return json.dumps([_transport_row_to_dict(r) for r in rows])
    return json.dumps([_row_to_dict(r) for r in rows])


def _menu_row_to_dict(row: ParsedMenuRow) -> dict:
    """Convert a ParsedMenuRow to a serializable dict."""
    return {
        "restaurant_name": row.restaurant_name or "",
        "city": row.city or "",
        "menu_name": row.menu_name,
        "description": row.description,
        "lunch_price": float(row.lunch_price) if row.lunch_price is not None else None,
        "dinner_price": float(row.dinner_price) if row.dinner_price is not None else None,
        "lunch_child_price": float(row.lunch_child_price) if row.lunch_child_price is not None else None,
        "dinner_child_price": float(row.dinner_child_price) if row.dinner_child_price is not None else None,
        "course_1": row.course_1,
        "course_2": row.course_2,
        "course_3": row.course_3,
        "course_4": row.course_4,
        "course_5": row.course_5,
        "min_pax": row.min_pax,
        "drink_included": row.drink_included,
        "season_code": row.season_code,
        "date_ranges": [
            {
                "date_from": dr.date_from.isoformat() if dr.date_from else None,
                "date_to": dr.date_to.isoformat() if dr.date_to else None,
            }
            for dr in row.date_ranges
        ],
        "note": row.note,
        "address": row.address,
        "phone": row.phone,
        "email": row.email,
    }


def _menu_row_dedup_key(row: ParsedMenuRow) -> tuple:
    """Build a hashable key from all meaningful fields of a menu row."""
    date_ranges_key = tuple(
        (dr.date_from, dr.date_to) for dr in row.date_ranges
    )
    return (
        (row.restaurant_name or "").strip().upper(),
        (row.city or "").strip().upper(),
        (row.menu_name or "").strip().upper(),
        row.lunch_price,
        row.dinner_price,
        row.lunch_child_price,
        row.dinner_child_price,
        row.min_pax,
        (row.season_code or "").strip().upper(),
        date_ranges_key,
    )


def deduplicate_menu_rows(rows: list[ParsedMenuRow]) -> list[ParsedMenuRow]:
    """Remove exact duplicate menu rows, preserving order."""
    seen: set = set()
    unique: list[ParsedMenuRow] = []
    for row in rows:
        key = _menu_row_dedup_key(row)
        if key not in seen:
            seen.add(key)
            unique.append(row)
    if len(unique) < len(rows):
        logger.info(f"Menu dedup removed {len(rows) - len(unique)} duplicate rows ({len(rows)} → {len(unique)})")
    return unique


def _transport_row_dedup_key(row: ParsedTransportRow) -> tuple:
    """Build a hashable key from all meaningful fields of a transport row."""
    return (
        (row.code or "").strip().upper(),
        row.price,
        (row.company_code or "").strip().upper(),
        (row.product or "").strip().upper(),
        row.bus_size,
        (row.service_type or "").strip().upper(),
        row.days,
        (row.route_description or "").strip().upper(),
        (row.city or "").strip().upper(),
    )


def deduplicate_transport_rows(rows: list[ParsedTransportRow]) -> list[ParsedTransportRow]:
    """Remove exact duplicate transport rows, preserving order."""
    seen: set = set()
    unique: list[ParsedTransportRow] = []
    for row in rows:
        key = _transport_row_dedup_key(row)
        if key not in seen:
            seen.add(key)
            unique.append(row)
    if len(unique) < len(rows):
        logger.info(f"Transport dedup removed {len(rows) - len(unique)} duplicate rows ({len(rows)} → {len(unique)})")
    return unique


def parse_document_background(document_id: int) -> None:
    """Parse a document in the background and cache the result.

    Opens its own DB session since the request session is already closed.
    """
    from app.database import SessionLocal
    from app.config import settings
    from app.services.confidence import score_extraction

    db = SessionLocal()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            logger.error(f"Background parse: document {document_id} not found")
            return

        file_path = settings.upload_path / document.filename
        if not file_path.exists():
            document.status = "failed"
            document.notes = "Source file not found on disk"
            db.commit()
            return

        # For restaurant/transport docs, parse_document handles parsed_rows_json internally
        rows = parse_document(file_path, document, db)
        if document.document_category not in ("restaurant", "transportation"):
            document.parsed_rows_json = rows_to_json(rows)
        document.status = "pending_review"
        document.row_count = document.row_count or len(rows)

        # Compute confidence score (Phase 2.1)
        if rows:
            document.confidence_score = score_extraction(rows)

        # Auto-assign to folder based on parsed data
        try:
            from app.services.folder_service import auto_assign_folder
            auto_assign_folder(document, db)
        except Exception as e:
            logger.warning("Auto-folder assignment failed for doc %d: %s", document_id, e)

        db.commit()
    except Exception as e:
        logger.exception(f"Background parse failed for document {document_id}")
        try:
            document.status = "failed"
            document.notes = str(e)[:500]
            db.commit()
        except Exception:
            db.rollback()
    finally:
        db.close()
