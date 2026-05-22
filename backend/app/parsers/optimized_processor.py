"""Optimized document processing pipeline.

Reduces AI costs by 80%+ through:
- Smart page truncation (send only relevant pages to AI)
- Haiku-first extraction (10x cheaper than Sonnet)
- Redis caching (no re-processing identical documents)
- Confidence-based routing (rule-based → Haiku → Sonnet fallback)
"""
from __future__ import annotations

import hashlib
import json
import logging
import re
import time
from dataclasses import dataclass, field
from pathlib import Path

from app.parsers.base import ParsedPriceRow
from app.parsers.category_detector import (
    HOTEL_KEYWORDS,
    RESTAURANT_KEYWORDS,
    TRANSPORTATION_KEYWORDS,
)

logger = logging.getLogger(__name__)

# Build keyword list from existing category keywords + pricing-specific terms
RELEVANT_KEYWORDS = list(set(
    HOTEL_KEYWORDS + RESTAURANT_KEYWORDS + TRANSPORTATION_KEYWORDS
)) + [
    "rate", "price", "eur", "usd", "€", "tax", "vat",
    "cancellation", "season", "commission", "allotment", "tarif",
]

# Regex for detecting price patterns on a page
PRICE_PATTERN = re.compile(r"\d+[\.,]?\d*\s*(?:EUR|USD|MAD|€|\$|DH)", re.IGNORECASE)


# ============================================================
#  Module 1: DocumentTruncator
# ============================================================

class DocumentTruncator:
    """Selects only relevant pages from a PDF before sending to AI."""

    def __init__(self, max_pages: int = 15):
        self.max_pages = max_pages

    def find_relevant_pages(self, file_path: str | Path) -> list[int]:
        """Scan pages for keywords, tables, and currency patterns.

        Returns list of 0-based page indices.
        """
        import pdfplumber

        file_path = Path(file_path)
        relevant: list[int] = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = (page.extract_text() or "").lower()
                    tables = page.extract_tables() or []

                    # Score criteria
                    keyword_matches = sum(1 for kw in RELEVANT_KEYWORDS if kw in text)
                    has_table = len(tables) > 0
                    has_prices = bool(PRICE_PATTERN.search(text))

                    if keyword_matches >= 2 or has_table or has_prices:
                        relevant.append(i)
        except Exception as e:
            logger.warning("Page relevance scan failed: %s", e)

        # Fallback: if no relevant pages found, take first max_pages
        if not relevant:
            try:
                with pdfplumber.open(file_path) as pdf:
                    relevant = list(range(min(len(pdf.pages), self.max_pages)))
            except Exception:
                relevant = list(range(self.max_pages))

        # Cap at max_pages
        return relevant[: self.max_pages]

    def extract_truncated_text(self, file_path: str | Path, relevant_pages: list[int]) -> str:
        """Extract text only from selected pages with [Page N] markers."""
        import pdfplumber

        file_path = Path(file_path)
        text_parts: list[str] = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_idx in relevant_pages:
                    if page_idx >= len(pdf.pages):
                        continue
                    page = pdf.pages[page_idx]
                    page_text = page.extract_text() or ""

                    # Also include table data
                    for table in page.extract_tables() or []:
                        for row in table:
                            if row:
                                cells = [str(c) for c in row if c]
                                if cells:
                                    page_text += "\n" + " | ".join(cells)

                    if page_text.strip():
                        text_parts.append(f"[Page {page_idx + 1}]\n{page_text}")
        except Exception as e:
            logger.warning("Truncated text extraction failed: %s", e)

        return "\n\n".join(text_parts)

    def get_truncation_hash(self, file_path: str | Path) -> str:
        """SHA-256 hash of truncated content for cache key."""
        relevant = self.find_relevant_pages(file_path)
        text = self.extract_truncated_text(file_path, relevant)
        return hashlib.sha256(text.encode("utf-8")).hexdigest()


# ============================================================
#  Module 2: HaikuExtractor
# ============================================================

HAIKU_EXTRACTION_PROMPT = """You are a hotel contract data extraction engine.

Extract ALL rate data from this document into a JSON array.

RULES:
1. One row per hotel × room_type × season × meal_plan combination
2. Normalize dates to YYYY-MM-DD format
3. Convert per-person rates to per-room:
   - DBL/TWN = per_person_rate × 2
   - SGL = per_person_rate + single_supplement
4. Apply single supplements inline to single_price
5. Meal plans: BB, HB, FB, AI, RO
6. Season codes: L (low), H (high), M (mid), P (peak), Annual
7. Strip currency symbols from prices
8. Stars as integer (1-5) or null
9. room_desc should be the room type name (Std, Superior, Suite, etc.)
   Never use "Single"/"Double" as room_desc — those are occupancy types
10. Set fit_git to "FIT" or "GIT"
11. Do NOT add taxes to prices — note them separately
12. Only create meal plan rows if the document has EXPLICIT rates for that meal plan

RESPOND ONLY with a valid JSON array. No backticks. No explanation.

Schema:
{{
  "accommodation": "Hotel name",
  "city": "City name",
  "room_desc": "Room type (Std, Superior, Suite, etc.)",
  "stars": integer or null,
  "hotel_type": "Hotel / Riad / Kasbah or null",
  "meal_plan": "BB / HB / FB / AI / RO or null",
  "fit_git": "FIT / GIT or null",
  "season_code": "L / H / M / P / Annual",
  "date_ranges": [{{"date_from": "YYYY-MM-DD", "date_to": "YYYY-MM-DD"}}],
  "double_price": number or null,
  "single_price": number or null,
  "twin_price": number or null,
  "triple_price": number or null,
  "quadruple_price": number or null,
  "baby_discount": "string or null",
  "child_discount": "string or null",
  "min_stay": integer or null,
  "note": "conditions, supplements, conversion notes",
  "address": "string or null",
  "phone": "string or null",
  "email": "string or null",
  "confidence": 0.0 to 1.0
}}"""


class HaikuExtractor:
    """Extract hotel contract data using Claude via Agent SDK."""

    MAX_CHARS = 30000

    def __init__(self, max_chars: int = 30000):
        self.max_chars = max_chars

    def extract(self, text: str) -> dict:
        """Extract data from text using Claude Agent SDK.

        Returns dict with:
        - rows: list[dict] — extracted row data
        - confidence: float — average confidence across rows
        """
        from app.services.claude_sdk import call_claude_sdk

        # Truncate text
        if len(text) > self.max_chars:
            text = text[: self.max_chars] + "\n... (truncated)"

        try:
            raw = call_claude_sdk(
                prompt=text,
                system_prompt=HAIKU_EXTRACTION_PROMPT,
            )
        except Exception as e:
            logger.error("Haiku extraction failed: %s", e)
            return {"rows": [], "confidence": 0.0}

        data = self._parse_json(raw)
        if not data:
            return {"rows": [], "confidence": 0.0}

        # Calculate average confidence
        confidences = [r.pop("confidence", 0.5) for r in data]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        return {"rows": data, "confidence": avg_confidence}

    @staticmethod
    def _parse_json(raw: str) -> list[dict]:
        """Parse JSON array from response, handling markdown fences and truncation."""
        text = raw.strip()
        if text.startswith("```"):
            lines = text.split("\n")
            end = len(lines) - 1
            while end > 0 and not lines[end].strip().startswith("```"):
                end -= 1
            if end > 0:
                text = "\n".join(lines[1:end])
            else:
                text = "\n".join(lines[1:])

        try:
            data = json.loads(text)
        except json.JSONDecodeError:
            # Try to recover truncated JSON
            last_complete = text.rfind("},")
            if last_complete == -1:
                last_complete = text.rfind("}\n]")
            if last_complete > 0:
                truncated = text[: last_complete + 1] + "\n]"
                try:
                    data = json.loads(truncated)
                    logger.info("Recovered %d rows from truncated Haiku JSON", len(data) if isinstance(data, list) else 0)
                except json.JSONDecodeError:
                    logger.warning("Haiku returned invalid JSON (recovery failed)")
                    return []
            else:
                logger.warning("Haiku returned invalid JSON, no recovery possible")
                return []

        if not isinstance(data, list):
            return []
        return [item for item in data if isinstance(item, dict)]


# ============================================================
#  Module 3: ExtractionCache (Redis)
# ============================================================

class ExtractionCache:
    """Redis-backed cache for extraction results."""

    KEY_PREFIX = "extraction:"

    def __init__(self, redis_url: str = "redis://localhost:6379", ttl_days: int = 30):
        self.ttl_seconds = ttl_days * 86400
        self._client = None
        self._available = False
        try:
            import redis
            self._client = redis.from_url(redis_url, decode_responses=True)
            self._client.ping()
            self._available = True
            logger.info("Redis cache connected: %s", redis_url)
        except Exception as e:
            logger.warning("Redis unavailable, caching disabled: %s", e)

    def get(self, document_hash: str) -> dict | None:
        """Get cached extraction result by document hash."""
        if not self._available:
            return None
        try:
            key = f"{self.KEY_PREFIX}{document_hash}"
            data = self._client.get(key)
            if data:
                return json.loads(data)
        except Exception as e:
            logger.warning("Redis GET failed: %s", e)
        return None

    def set(self, document_hash: str, result: dict) -> None:
        """Cache extraction result with TTL."""
        if not self._available:
            return
        try:
            key = f"{self.KEY_PREFIX}{document_hash}"
            self._client.setex(key, self.ttl_seconds, json.dumps(result, default=str))
        except Exception as e:
            logger.warning("Redis SET failed: %s", e)


# ============================================================
#  Module 4: CostMetrics
# ============================================================

@dataclass
class CostMetrics:
    """Track AI usage costs for the pipeline."""

    haiku_calls: int = 0
    sonnet_calls: int = 0
    rule_based_hits: int = 0
    cache_hits: int = 0
    hybrid_mode_used: int = 0

    HAIKU_COST = 0.01
    SONNET_COST = 0.15

    @property
    def estimated_cost(self) -> float:
        return (self.haiku_calls * self.HAIKU_COST) + (self.sonnet_calls * self.SONNET_COST)

    def log_summary(self) -> None:
        logger.info(
            "Pipeline cost summary — "
            "cache_hits=%d, rule_based=%d, haiku=%d, sonnet=%d, hybrid=%d, "
            "estimated_cost=$%.3f",
            self.cache_hits,
            self.rule_based_hits,
            self.haiku_calls,
            self.sonnet_calls,
            self.hybrid_mode_used,
            self.estimated_cost,
        )


# ============================================================
#  Module 5: OptimizedDocumentProcessor
# ============================================================

class OptimizedDocumentProcessor:
    """Orchestrates the optimized extraction pipeline.

    Flow:
    1. Check Redis cache → return if hit
    2. Truncate PDF to relevant pages
    3. Try rule-based parser → quality score
    4. quality >= 0.8: use rule-based ($0)
    5. quality >= 0.5: hybrid (rule-based + Haiku fill gaps, ~$0.01)
    6. quality < 0.5: full Haiku (~$0.01)
    7. Haiku confidence < validation_threshold: Sonnet fallback
    8. Cache result in Redis
    """

    def __init__(
        self,
        anthropic_api_key: str = "",
        redis_url: str = "redis://localhost:6379",
        max_pages: int = 15,
        max_chars: int = 30000,
        enable_sonnet_fallback: bool = True,
        validation_threshold: float = 0.6,
        cache_ttl_days: int = 30,
    ):
        self.max_pages = max_pages
        self.max_chars = max_chars
        self.enable_sonnet_fallback = enable_sonnet_fallback
        self.validation_threshold = validation_threshold

        self.truncator = DocumentTruncator(max_pages=max_pages)
        self.haiku = HaikuExtractor(max_chars=max_chars)
        self.cache = ExtractionCache(redis_url=redis_url, ttl_days=cache_ttl_days)
        self.metrics = CostMetrics()

    def process_document(
        self,
        file_path: str | Path,
        document_hash: str | None = None,
    ) -> dict:
        """Main entry point for optimized extraction.

        Returns dict with:
        - rows: list[ParsedPriceRow]
        - confidence: float
        - extraction_method: str
        - metrics: CostMetrics
        """
        file_path = Path(file_path)

        # 1. Check Redis cache
        cache_key = document_hash
        if cache_key:
            cached = self.cache.get(cache_key)
            if cached:
                self.metrics.cache_hits += 1
                logger.info("Cache hit for document hash %s", cache_key[:12])
                # Convert cached dicts back to ParsedPriceRow
                rows = self._dicts_to_rows(cached.get("rows", []))
                return {
                    "rows": rows,
                    "confidence": cached.get("confidence", 1.0),
                    "extraction_method": "cache",
                    "metrics": self.metrics,
                }

        # 2. Truncate PDF to relevant pages (for text extraction)
        file_type = file_path.suffix.lower().lstrip(".")
        truncated_text = None
        relevant_pages = None

        if file_type == "pdf":
            relevant_pages = self.truncator.find_relevant_pages(file_path)
            truncated_text = self.truncator.extract_truncated_text(file_path, relevant_pages)
            logger.info(
                "Truncated PDF to %d/%s relevant pages",
                len(relevant_pages),
                "?" if relevant_pages else "0",
            )

        # 3. Try rule-based parser → quality score
        rule_rows, quality = self._try_rule_based(file_path, file_type)
        logger.info("Rule-based quality: %.3f (%d rows)", quality, len(rule_rows))

        # 4-6. Quality-based routing
        rows: list[ParsedPriceRow] = []
        confidence = quality
        method = "rule_based"

        if quality >= 0.8:
            # High quality — trust rule-based
            rows = rule_rows
            confidence = quality
            method = "rule_based"
            self.metrics.rule_based_hits += 1

        elif quality >= 0.5:
            # Medium quality — hybrid: rule-based + Haiku fills gaps
            text = truncated_text or self._extract_text(file_path, file_type)
            if text:
                haiku_result = self.haiku.extract(text)
                self.metrics.haiku_calls += 1
                self.metrics.hybrid_mode_used += 1

                haiku_rows = self._dicts_to_rows(haiku_result["rows"])
                rows = self._merge_results(rule_rows, haiku_rows)
                confidence = max(quality, haiku_result["confidence"])
                method = "hybrid"
                logger.info(
                    "Hybrid: %d rule + %d haiku → %d merged (confidence=%.2f)",
                    len(rule_rows), len(haiku_rows), len(rows), confidence,
                )
            else:
                rows = rule_rows
                method = "rule_based"
                self.metrics.rule_based_hits += 1

        else:
            # Low quality — full Haiku extraction
            text = truncated_text or self._extract_text(file_path, file_type)
            if text:
                haiku_result = self.haiku.extract(text)
                self.metrics.haiku_calls += 1
                rows = self._dicts_to_rows(haiku_result["rows"])
                confidence = haiku_result["confidence"]
                method = "haiku"
                logger.info("Haiku extraction: %d rows (confidence=%.2f)", len(rows), confidence)
            else:
                rows = rule_rows
                method = "rule_based"

        # 7. Sonnet fallback if confidence is too low
        if (
            confidence < self.validation_threshold
            and self.enable_sonnet_fallback
        ):
            logger.info(
                "Confidence %.2f < threshold %.2f — falling back to Sonnet",
                confidence, self.validation_threshold,
            )
            sonnet_rows = self._sonnet_fallback(file_path, file_type, truncated_text)
            if sonnet_rows:
                rows = sonnet_rows
                confidence = 0.85  # Sonnet is our most reliable
                method = "sonnet_fallback"
                self.metrics.sonnet_calls += 1

        # 8. Cache result in Redis
        if cache_key and rows:
            cache_data = {
                "rows": self._rows_to_dicts(rows),
                "confidence": confidence,
                "extraction_method": method,
            }
            self.cache.set(cache_key, cache_data)

        self.metrics.log_summary()

        return {
            "rows": rows,
            "confidence": confidence,
            "extraction_method": method,
            "metrics": self.metrics,
        }

    def _try_rule_based(self, file_path: Path, file_type: str) -> tuple[list[ParsedPriceRow], float]:
        """Run rule-based parser and score the result."""
        from app.parsers.excel_parser import ExcelParser, CsvParser
        from app.parsers.pdf_parser import PdfParser
        from app.parsers.word_parser import WordParser
        from app.parsers.text_parser import TextParser
        from app.parsers.email_parser import EmailParser

        parsers = {
            "xlsx": ExcelParser(),
            "xls": ExcelParser(),
            "csv": CsvParser(),
            "pdf": PdfParser(),
            "docx": WordParser(),
            "doc": WordParser(),
            "txt": TextParser(),
            "text": TextParser(),
            "eml": EmailParser(),
            "msg": EmailParser(),
        }

        parser = parsers.get(file_type)
        if not parser:
            return [], 0.0

        try:
            rows = parser.parse(file_path)
        except Exception as e:
            logger.warning("Rule-based parser failed: %s", e)
            return [], 0.0

        quality = self._quality_score(rows)
        return rows, quality

    @staticmethod
    def _quality_score(rows: list[ParsedPriceRow]) -> float:
        """Score quality of rule-based extraction (0.0 – 1.0)."""
        if not rows:
            return 0.0

        n = len(rows)
        has_price = sum(
            1 for r in rows
            if any([r.double_price, r.single_price, r.twin_price, r.triple_price, r.quadruple_price])
        )
        has_name = sum(1 for r in rows if (r.accommodation or "").strip())
        has_dates = sum(1 for r in rows if r.date_ranges)

        score = (
            0.5 * (has_price / n)
            + 0.3 * (has_name / n)
            + 0.2 * (has_dates / n)
        )
        return round(score, 3)

    def _extract_text(self, file_path: Path, file_type: str) -> str | None:
        """Extract text from non-PDF documents."""
        if file_type in ("docx", "doc"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        text_parts.append(" | ".join(cells))
            return "\n".join(text_parts)
        elif file_type in ("txt", "csv", "text"):
            return file_path.read_text(encoding="utf-8", errors="replace")
        elif file_type in ("xlsx", "xls"):
            import pandas as pd
            dfs = pd.read_excel(file_path, sheet_name=None, header=None)
            text_parts = []
            for sheet_name, df in dfs.items():
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                for _, row in df.iterrows():
                    cells = [str(c) if pd.notna(c) else "" for c in row]
                    if any(cells):
                        text_parts.append(" | ".join(cells))
            return "\n".join(text_parts)
        elif file_type in ("eml", "msg"):
            return self._extract_email_text(file_path, file_type)
        return None

    @staticmethod
    def _extract_email_text(file_path: Path, file_type: str) -> str | None:
        """Extract body text and attachment text from an email file."""
        text_parts: list[str] = []

        if file_type == "msg":
            try:
                import extract_msg
                msg = extract_msg.Message(str(file_path))
                if msg.subject:
                    text_parts.append(f"Subject: {msg.subject}")
                if msg.sender:
                    text_parts.append(f"From: {msg.sender}")
                if msg.body:
                    text_parts.append(msg.body)
            except Exception:
                pass
        else:
            try:
                import email as email_lib
                from email import policy as email_policy
                raw = file_path.read_bytes()
                msg = email_lib.message_from_bytes(raw, policy=email_policy.default)
                subject = msg.get("Subject", "")
                sender = msg.get("From", "")
                if subject:
                    text_parts.append(f"Subject: {subject}")
                if sender:
                    text_parts.append(f"From: {sender}")
                body = msg.get_body(preferencelist=("plain", "html"))
                if body:
                    content = body.get_content()
                    if isinstance(content, str):
                        text_parts.append(content)
                # Extract text from non-binary attachments
                for part in msg.walk():
                    disposition = str(part.get("Content-Disposition", ""))
                    if "attachment" not in disposition:
                        continue
                    fname = part.get_filename()
                    if not fname:
                        continue
                    from app.utils import detect_file_type
                    att_type = detect_file_type(fname)
                    if att_type in ("txt", "csv", "text"):
                        payload = part.get_payload(decode=True)
                        if payload:
                            text_parts.append(f"--- Attachment: {fname} ---")
                            text_parts.append(payload.decode("utf-8", errors="replace"))
            except Exception:
                pass

        return "\n".join(text_parts) if text_parts else None

    def _sonnet_fallback(
        self, file_path: Path, file_type: str, truncated_text: str | None
    ) -> list[ParsedPriceRow]:
        """Fall back to full Sonnet extraction (existing AiParser)."""
        from app.parsers.ai_parser import AiParser

        ai_parser = AiParser()
        try:
            if file_type == "pdf":
                return ai_parser.parse_pdf(file_path)
            else:
                text = truncated_text or self._extract_text(file_path, file_type)
                if text:
                    return ai_parser.parse_text(text)
        except Exception as e:
            logger.error("Sonnet fallback failed: %s", e)
        return []

    def _merge_results(
        self,
        primary: list[ParsedPriceRow],
        secondary: list[ParsedPriceRow],
    ) -> list[ParsedPriceRow]:
        """Merge primary (rule-based) and secondary (AI) rows.

        Uses (accommodation, room_desc, season_code, meal_plan) as merge key.
        Primary rows are kept; secondary fills in nulls.
        Additional secondary rows not in primary are appended.
        """
        from dataclasses import replace

        def row_key(r: ParsedPriceRow) -> tuple:
            return (
                (r.accommodation or "").strip().upper(),
                (r.room_desc or "").strip().upper(),
                (r.season_code or "").strip().upper(),
                (r.meal_plan or "").strip().upper(),
            )

        # Index primary rows by key
        primary_map: dict[tuple, ParsedPriceRow] = {}
        for r in primary:
            key = row_key(r)
            primary_map[key] = r

        merged: list[ParsedPriceRow] = []
        secondary_keys_used: set[tuple] = set()

        for key, prow in primary_map.items():
            # Find matching secondary row
            match = None
            for srow in secondary:
                if row_key(srow) == key:
                    match = srow
                    secondary_keys_used.add(key)
                    break

            if match:
                # Fill null fields in primary from secondary
                updates = {}
                for attr in (
                    "double_price", "single_price", "twin_price",
                    "triple_price", "quadruple_price", "stars",
                    "hotel_type", "meal_plan", "fit_git", "baby_discount",
                    "child_discount", "min_stay", "note", "address",
                    "phone", "email",
                ):
                    if getattr(prow, attr) is None and getattr(match, attr) is not None:
                        updates[attr] = getattr(match, attr)
                # Fill date_ranges if primary has none
                if not prow.date_ranges and match.date_ranges:
                    updates["date_ranges"] = match.date_ranges

                if updates:
                    prow = replace(prow, **updates)

            merged.append(prow)

        # Append any secondary rows not matched
        for srow in secondary:
            if row_key(srow) not in secondary_keys_used:
                merged.append(srow)

        return merged

    def _dicts_to_rows(self, data: list[dict]) -> list[ParsedPriceRow]:
        """Convert raw dicts to ParsedPriceRow using AiParser._convert_rows."""
        from app.parsers.ai_parser import AiParser
        parser = AiParser()
        return parser._convert_rows(data)

    @staticmethod
    def _rows_to_dicts(rows: list[ParsedPriceRow]) -> list[dict]:
        """Convert ParsedPriceRow objects to dicts for caching."""
        from app.services.extraction import _row_to_dict
        return [_row_to_dict(r) for r in rows]
