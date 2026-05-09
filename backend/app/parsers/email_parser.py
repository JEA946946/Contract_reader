from __future__ import annotations

import email
from email import policy
from pathlib import Path
import tempfile

from app.parsers.base import BaseParser, ParsedPriceRow
from app.parsers.text_parser import TextParser
from app.parsers.excel_parser import ExcelParser, CsvParser
from app.parsers.pdf_parser import PdfParser
from app.parsers.word_parser import WordParser
from app.utils import detect_file_type


class EmailParser(BaseParser):
    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in (".eml", ".msg")

    def parse(self, file_path: Path) -> list[ParsedPriceRow]:
        if file_path.suffix.lower() == ".msg":
            return self._parse_msg(file_path)
        return self._parse_eml(file_path)

    def _parse_eml(self, file_path: Path) -> list[ParsedPriceRow]:
        all_rows = []

        with open(file_path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=policy.default)

        # Build email context for AI parser
        email_context = _build_email_context(
            subject=msg.get("Subject", ""),
            sender=msg.get("From", ""),
            body=msg.get_body(preferencelist=("plain",)),
        )

        # Extract body text
        body = msg.get_body(preferencelist=("plain", "html"))
        if body:
            text_content = body.get_content()
            if text_content:
                text_parser = TextParser()
                temp = file_path.with_suffix(".tmp.txt")
                temp.write_text(text_content)
                try:
                    rows = text_parser.parse(temp)
                    all_rows.extend(rows)
                finally:
                    temp.unlink(missing_ok=True)

        # Process attachments — use unique temp files to avoid overwriting
        # the original .eml when an attachment is itself an .eml
        att_idx = 0
        for part in msg.walk():
            if part.get_content_disposition() in ("attachment", "inline") and part.get_filename():
                filename = part.get_filename()
                if not filename:
                    continue
                data = part.get_payload(decode=True)
                if not data:
                    continue

                suffix = Path(filename).suffix
                temp_file = file_path.parent / f"{file_path.stem}_att{att_idx}{suffix}"
                att_idx += 1
                temp_file.write_bytes(data)
                try:
                    file_type = detect_file_type(filename)
                    rows = _parse_attachment(temp_file, file_type, email_context)
                    all_rows.extend(rows)
                finally:
                    temp_file.unlink(missing_ok=True)

        return all_rows

    def _parse_msg(self, file_path: Path) -> list[ParsedPriceRow]:
        import extract_msg

        all_rows = []
        msg = extract_msg.Message(str(file_path))

        email_context = _build_email_context(
            subject=msg.subject or "",
            sender=msg.sender or "",
            body_text=msg.body,
        )

        if msg.body:
            text_parser = TextParser()
            temp = file_path.with_suffix(".tmp.txt")
            temp.write_text(msg.body)
            try:
                rows = text_parser.parse(temp)
                all_rows.extend(rows)
            finally:
                temp.unlink(missing_ok=True)

        for att_idx, attachment in enumerate(msg.attachments):
            if not attachment.longFilename:
                continue
            suffix = Path(attachment.longFilename).suffix
            temp_file = file_path.parent / f"{file_path.stem}_att{att_idx}{suffix}"
            temp_file.write_bytes(attachment.data)
            try:
                file_type = detect_file_type(attachment.longFilename)
                rows = _parse_attachment(temp_file, file_type, email_context)
                all_rows.extend(rows)
            finally:
                temp_file.unlink(missing_ok=True)

        return all_rows


def _build_email_context(subject: str = "", sender: str = "", body=None, body_text: str | None = None) -> str:
    """Build a context string from email metadata to help AI identify hotel/restaurant."""
    parts = []
    if sender:
        parts.append(f"Email sender: {sender}")
    if subject:
        parts.append(f"Email subject: {subject}")
    if body_text:
        snippet = body_text[:500]
    elif body:
        try:
            snippet = body.get_content()[:500]
        except Exception:
            snippet = ""
    else:
        snippet = ""
    if snippet:
        parts.append(f"Email body excerpt:\n{snippet}")
    return "\n".join(parts)


def _get_parser_for_type(file_type: str) -> BaseParser | None:
    parsers = {
        "xlsx": ExcelParser(),
        "xls": ExcelParser(),
        "csv": CsvParser(),
        "pdf": PdfParser(),
        "docx": WordParser(),
        "txt": TextParser(),
    }
    return parsers.get(file_type)


def _parse_attachment(temp_file: Path, file_type: str, email_context: str = "") -> list[ParsedPriceRow]:
    """Parse an email attachment. PDFs always use AI vision; others try rule-based first.

    email_context: metadata from the parent email (sender, subject, body snippet)
    to help the AI parser identify hotel/restaurant names.
    """
    rows: list[ParsedPriceRow] = []

    context_header = ""
    if email_context:
        context_header = (
            "--- EMAIL CONTEXT (use this to identify the hotel/property name and city) ---\n"
            f"{email_context}\n"
            "--- END EMAIL CONTEXT ---\n\n"
        )

    # For PDFs, always use AI vision parser — rule-based is too unreliable
    # for complex hotel contract table layouts
    if file_type == "pdf":
        try:
            from app.parsers.ai_parser import AiParser
            rows = AiParser().parse_pdf(temp_file)
        except Exception:
            pass
        return rows

    # For other types, try rule-based first
    parser = _get_parser_for_type(file_type)
    if parser:
        try:
            rows = parser.parse(temp_file)
        except Exception:
            pass

    if not _has_meaningful_data(rows):
        rows = []
        try:
            from app.parsers.ai_parser import AiParser
            ai = AiParser()
            if file_type in ("docx", "doc"):
                from docx import Document as DocxDocument
                doc = DocxDocument(temp_file)
                text_parts = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            text_parts.append(" | ".join(cells))
                rows = ai.parse_text(context_header + "\n".join(text_parts))
            elif file_type in ("xlsx", "xls"):
                import pandas as pd
                dfs = pd.read_excel(temp_file, sheet_name=None, header=None)
                text_parts = []
                for sheet_name, df in dfs.items():
                    text_parts.append(f"--- Sheet: {sheet_name} ---")
                    for _, data_row in df.iterrows():
                        cells = [str(c) if pd.notna(c) else "" for c in data_row]
                        if any(cells):
                            text_parts.append(" | ".join(cells))
                rows = ai.parse_text(context_header + "\n".join(text_parts))
            elif file_type in ("txt", "csv", "text"):
                if context_header:
                    text = context_header + temp_file.read_text(encoding="utf-8", errors="replace")
                    rows = ai.parse_text(text)
                else:
                    rows = ai.parse(temp_file)
        except Exception:
            pass

    return rows


def _has_meaningful_data(rows: list[ParsedPriceRow]) -> bool:
    """Check if any row has at least one price value."""
    return any(r.double_price or r.single_price or r.twin_price or r.triple_price or r.quadruple_price for r in rows)
