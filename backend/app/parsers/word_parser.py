from __future__ import annotations

import re
from docx import Document as DocxDocument
import pandas as pd
from pathlib import Path

from app.parsers.base import BaseParser, ParsedPriceRow, ParsedDateRange
from app.parsers.excel_parser import ExcelParser
from app.utils import parse_price, clean_string

# Pattern to match lines like "Chambre Double ………… 600 Dhs / TTC"
# or "Chambre Single / Double / Triple en D.P ... : 350Dhs/ Pax"
PRICE_LINE_PATTERN = re.compile(
    r'^(.+?)\s*[…\.,:]+\s*:?\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b',
    re.IGNORECASE
)

# Pattern for "- Chambre Single : 675 DH /595 DH" style (dual price DP/BB)
DASH_PRICE_PATTERN = re.compile(
    r'^[-–•]\s*(.+?)\s*:\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b'
    r'(?:\s*/\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b)?',
    re.IGNORECASE
)

# Pattern for "Lit Sup en DP : 475 DH" style (single price with dash prefix)
EXTRA_PRICE_PATTERN = re.compile(
    r'^[-–•]?\s*(.+?)\s*:\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b',
    re.IGNORECASE
)

# Patterns to detect room types
ROOM_PATTERNS = {
    'suite_royale': re.compile(r'suite\s+royale', re.IGNORECASE),
    'suite_junior': re.compile(r'suite\s+junior', re.IGNORECASE),
    'suite': re.compile(r'suite', re.IGNORECASE),
    'single': re.compile(r'single', re.IGNORECASE),
    'double': re.compile(r'double', re.IGNORECASE),
    'triple': re.compile(r'triple', re.IGNORECASE),
    'chambre': re.compile(r'chambre', re.IGNORECASE),
}

# Meal plan patterns
MEAL_PATTERNS = {
    'DP': re.compile(r'\bD\.?P\b', re.IGNORECASE),
    'BB': re.compile(r'\bB\.?B\b', re.IGNORECASE),
    'AI': re.compile(r'\bAll\s*In\b', re.IGNORECASE),
    'FB': re.compile(r'\bF\.?B\b', re.IGNORECASE),
    'HB': re.compile(r'\bH\.?B\b', re.IGNORECASE),
}


def extract_hotel_info_from_doc(doc):
    """Extract hotel name, city, stars, address, phone, and email from document headers and paragraphs."""
    hotel_name = "Unknown Hotel"
    city = ""
    stars = None
    address = ""
    phone = ""
    email_addr = ""

    # First: check document headers (most reliable source)
    header_text_all = ""
    for section in doc.sections:
        header = section.header
        if header and header.paragraphs:
            for p in header.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                header_text_all += " " + text

                # Hotel name
                hotel_match = re.search(
                    r'((?:h[oô]tel|riad|dar|kasbah|maison)\s+[\w\s\'\-]+?)(?:\s*[★\*]|$|\s+\d\s*[★\*éeE]|\s+adresse|\s+tel|\s+fax)',
                    text, re.IGNORECASE
                )
                if hotel_match and hotel_name == "Unknown Hotel":
                    hotel_name = hotel_match.group(1).strip()
                    # Stars as icons: ★★★★ or ****
                    stars_match = re.search(r'([★]{1,5}|\*{1,5})', text)
                    if stars_match:
                        stars = len(stars_match.group(1))
                    else:
                        # Stars as number: "4*", "4 étoiles", "4 stars"
                        num_match = re.search(r'([1-5])\s*(?:[★\*]|[ée]toiles?|stars?)', text, re.IGNORECASE)
                        if num_match:
                            stars = int(num_match.group(1))

    # Extract address from header (after "Adresse :" or between hotel name and Tel)
    addr_match = re.search(
        r'[Aa]dresse\s*:\s*(.+?)(?:\s*[Tt]el|$)',
        header_text_all
    )
    if addr_match:
        address = addr_match.group(1).strip().rstrip('/')

    # Extract phone number (after "Tel :" up to "/" or "Fax")
    phone_match = re.search(
        r'[Tt]el\s*:\s*([\d\s\+]+?)(?:\s*/|\s*[Ff]ax|\s*$)',
        header_text_all
    )
    if phone_match:
        phone = phone_match.group(1).strip()
    # Also try to get fax
    fax_match = re.search(
        r'[Ff]ax\s*:\s*([\d\s\+]+?)(?:\s*/|\s*[Ee]|\s*$)',
        header_text_all
    )
    if fax_match:
        fax = fax_match.group(1).strip()
        if phone and fax:
            phone = f"Tel: {phone} / Fax: {fax}"
        elif fax:
            phone = f"Fax: {fax}"

    # Extract email
    email_match = re.search(
        r'[Ee][\-.]?[Mm]ail\s*:?\s*([\w\.\-]+@[\w\.\-]+\.\w+)',
        header_text_all
    )
    if email_match:
        email_addr = email_match.group(1)

    # Also check footer for contact info
    for section in doc.sections:
        footer = section.footer
        if footer and footer.paragraphs:
            for p in footer.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                # Email from footer
                if not email_addr:
                    em = re.search(r'([\w\.\-]+@[\w\.\-]+\.\w+)', text)
                    if em:
                        email_addr = em.group(1)
                # Phone from footer
                if not phone:
                    ph = re.search(r'[Tt]el\s*:?\s*([\d\s\+\-/]+)', text)
                    if ph:
                        phone = ph.group(1).strip()

    # Check body paragraphs for city
    for p in doc.paragraphs[:15]:
        text = p.text.strip()
        if not text:
            continue
        location_match = re.match(r'^(\w+)\s+[Ll]e\s+\d', text)
        if location_match:
            city = location_match.group(1)
            break

    # If no city found, try header address
    if not city:
        combined = header_text_all.lower()
        city_match = re.search(
            r'(?:midelt|marrakech|agadir|casablanca|fes|fez|rabat|tangier|tanger|'
            r'ouarzazate|essaouira|meknes|meknès|errachidia|ifrane|azrou|beni mellal|'
            r'nador|oujda|tetouan|chefchaouen|dakhla|kenitra|mohammedia|safi|el jadida|'
            r'afourer|tiznit|taroudant|zagora)',
            combined, re.IGNORECASE
        )
        if city_match:
            city = city_match.group(0).capitalize()

    return hotel_name, city, stars, address, phone, email_addr


def _extract_room_type(description: str) -> str:
    """Extract room type from a description string.

    Returns the room category (e.g. 'Suite Royale', 'Suite Junior', 'Chambre',
    'Mini Suite', 'Studio', 'Appartement', 'Villa', 'Bungalow').
    Falls back to 'Chambre' if no specific type is detected.
    """
    desc = description.lower()
    # Order matters: check most specific first
    if re.search(r'suite\s+royale', desc):
        return "Suite Royale"
    if re.search(r'suite\s+junior', desc):
        return "Suite Junior"
    if re.search(r'suite\s+sup[ée]rieure', desc):
        return "Suite Supérieure"
    if re.search(r'mini[\s-]suite', desc):
        return "Mini Suite"
    if re.search(r'suite\s+familiale', desc):
        return "Suite Familiale"
    if 'suite' in desc:
        return "Suite"
    if 'studio' in desc:
        return "Studio"
    if re.search(r'app?artement', desc):
        return "Appartement"
    if 'villa' in desc:
        return "Villa"
    if 'bungalow' in desc:
        return "Bungalow"
    if 'riad' in desc:
        return "Riad"
    if 'chambre' in desc or 'room' in desc:
        return "Chambre"
    return "Chambre"


MOROCCAN_CITIES = (
    'midelt|marrakech|agadir|casablanca|fes|fez|rabat|tangier|tanger|'
    'ouarzazate|essaouira|meknes|meknès|errachidia|ifrane|azrou|beni mellal|'
    'nador|oujda|tetouan|chefchaouen|dakhla|kenitra|mohammedia|safi|el jadida|'
    'afourer|tiznit|taroudant|zagora|merzouga|tinghir|todra|ait ben haddou|'
    'fnideq|al hoceima|larache|khouribga|settat|taza|guelmim'
)


def _extract_hotel_from_body(paragraphs, existing_city: str = ""):
    """Extract hotel name, city, stars, and phone from body paragraphs.

    Handles patterns like:
        CASA HASSAN / DAR BAIBOU
        22  RUE TARGUI
        Chefchaouen  (Morocco)
        Téléphone: (+212) 539-988196
        L'HOTEL COTE OCEAN MOGADOR 4*
    """
    hotel_name = "Unknown Hotel"
    city = existing_city
    phone = ""
    stars = None

    all_text = [p.text.strip() for p in paragraphs if p.text.strip()]

    for i, text in enumerate(all_text):
        # Look for hotel name patterns (CASA HASSAN, RIAD X, DAR X, etc.)
        # Standalone line: "HOTEL COTE OCEAN MOGADOR 4*" or "CASA HASSAN / DAR BAIBOU"
        # Also match "L'HOTEL X 4*" (with various apostrophe chars)
        name_match = re.match(
            r"^(?:l['\u2018\u2019\u0027])?\s*((?:CASA|RIAD|DAR|HOTEL|KASBAH|MAISON)\s+[\w\s'\-/]+?)(?:\s+(\d)\s*[★\*])?\s*$",
            text, re.IGNORECASE
        )
        if not name_match:
            # Embedded: "L'HOTEL COTE OCEAN MOGADOR 4*," or "Entre L'HOTEL X"
            name_match = re.search(
                r"(?:l['\u2018\u2019\u0027]?|entre\s+l['\u2018\u2019\u0027]?)\s*"
                r"((?:H[OÔ]TEL|RIAD|DAR|KASBAH|MAISON|CASA)\s+[\w\s'\-]+?)"
                r"(?:\s+(\d)\s*[★\*])?"
                r"(?:\s*[,;.]|\s*$)",
                text, re.IGNORECASE
            )
            # Reject sentence matches: require the word after the keyword to be uppercase
            # e.g. "HOTEL COTE OCEAN" OK, "hôtel attribuera les" rejected
            if name_match:
                matched_name = name_match.group(1).strip()
                name_words = matched_name.split()
                if len(name_words) >= 2 and not name_words[1][0].isupper():
                    name_match = None
        if name_match and hotel_name == "Unknown Hotel":
            hotel_name = name_match.group(1).strip()
            # Extract stars from "4*" pattern
            if name_match.group(2):
                stars = int(name_match.group(2))
            # Take the first part if there's a "/"
            if '/' in hotel_name:
                hotel_name = hotel_name.split('/')[0].strip()
            # Clean trailing star patterns that might have been captured in name
            hotel_name = re.sub(r'\s*\d\s*[★\*]\s*$', '', hotel_name).strip()

        # Look for city with "(Morocco)" or "(Maroc)"
        city_match = re.search(
            r'(' + MOROCCAN_CITIES + r')\s*(?:\(|morocco|maroc)?',
            text, re.IGNORECASE
        )
        if city_match and not city:
            city = city_match.group(1).strip().capitalize()

        # Phone
        phone_match = re.search(
            r'[Tt][ée]l[ée]phone\s*:?\s*([\d\s\+\(\)\-]+)',
            text
        )
        if phone_match and not phone:
            phone = phone_match.group(1).strip()

    return hotel_name, city, stars, phone


def _detect_pricing_mode(text: str):
    """Detect if price is per pax, per room for N pax, or per room.

    Returns:
        ("per_pax", 1)     — price is per person
        ("per_room", 2)    — price is per room for 2 pax (e.g. "/02Pax")
        ("per_room", 1)    — price is per room (e.g. "/ TTC" or no indicator)
    """
    # "/02Pax" or "/ 02 Pax" — per room for N people
    m = re.search(r'/\s*(\d{1,2})\s*Pax', text, re.IGNORECASE)
    if m:
        return "per_room", int(m.group(1))
    # "/Pax" or "/ Pax" — per person
    if re.search(r'/\s*Pax', text, re.IGNORECASE):
        return "per_pax", 1
    # Default: per room
    return "per_room", 1


def parse_unstructured_prices(paragraphs) -> list[ParsedPriceRow]:
    """Parse prices from unstructured paragraph text (dotted line format).

    Handles per-pax pricing: when a price is marked "/Pax", calculates
    room prices using the tax from the same section:
        double = 2 × (pax_price + tax)
        single = 1 × (pax_price + tax) + sup_single
        twin   = 2 × (pax_price + tax)
    """
    from decimal import Decimal

    # ── First pass: collect tax and sup-single per section ──
    section_extras: dict[str, dict] = {}
    current_category = ""

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        text_lower = text.lower()

        if 'tarifs public' in text_lower or 'tarif public' in text_lower:
            current_category = "Public"
            continue
        elif 'tarifs groupe' in text_lower or 'tarif groupe' in text_lower:
            current_category = "Groupe (GIT)"
            continue
        elif 'tarifs individuel' in text_lower or 'tarif individuel' in text_lower:
            current_category = "Individuel (FIT)"
            continue
        elif 'tarifs suite' in text_lower or 'tarif suite' in text_lower:
            current_category = "Suites"
            continue

        if current_category == "Public" or not current_category:
            continue

        match = PRICE_LINE_PATTERN.match(text)
        if not match:
            continue
        description = match.group(1).strip()
        price_str = match.group(2).strip()
        desc_lower = description.lower()

        if current_category not in section_extras:
            section_extras[current_category] = {"tax": Decimal("0"), "sup_single": Decimal("0")}

        if any(t in desc_lower for t in ['taxe', 't.p.t', 'tpt']):
            val = parse_price(price_str)
            if val:
                section_extras[current_category]["tax"] = val
        elif 'sup' in desc_lower and 'single' in desc_lower:
            val = parse_price(price_str)
            if val:
                section_extras[current_category]["sup_single"] = val

    # ── Second pass: build rows with calculated prices ──
    rows = []
    current_category = ""

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        text_lower = text.lower()

        if 'tarifs public' in text_lower or 'tarif public' in text_lower:
            current_category = "Public"
            continue
        elif 'tarifs groupe' in text_lower or 'tarif groupe' in text_lower:
            current_category = "Groupe (GIT)"
            continue
        elif 'tarifs individuel' in text_lower or 'tarif individuel' in text_lower:
            current_category = "Individuel (FIT)"
            continue
        elif 'tarifs suite' in text_lower or 'tarif suite' in text_lower:
            current_category = "Suites"
            continue

        if current_category == "Public" or not current_category:
            continue

        match = PRICE_LINE_PATTERN.match(text)
        if not match:
            continue

        description = match.group(1).strip()
        price_str = match.group(2).strip()
        desc_lower = description.lower()

        # Skip tax, sup-single, meals — already collected in first pass
        if any(skip in desc_lower for skip in [
            'taxe', 't.p.t', 't.c', 'tpt', 'petit', 'repas',
            'déjeuner', 'dejeuner', 'dîner', 'diner',
        ]):
            continue
        if 'sup' in desc_lower and 'single' in desc_lower:
            continue
        # Skip "1/2 chambre" lines — handled by parse_half_room_prices
        if '1/2' in desc_lower and 'chambre' in desc_lower:
            continue
        # Skip standalone supplement lines
        if desc_lower.startswith('suppl') or desc_lower.startswith('sup '):
            continue

        price = parse_price(price_str)
        if not price:
            continue

        # Detect meal plan
        meal_plan = ""
        for plan, pattern in MEAL_PATTERNS.items():
            if pattern.search(description):
                meal_plan = plan
                break

        # Detect room types mentioned
        is_single = 'single' in desc_lower
        is_double = 'double' in desc_lower
        is_triple = 'triple' in desc_lower
        is_suite = 'suite' in desc_lower

        # Determine FIT/GIT from category
        fit_git = None
        if 'groupe' in current_category.lower() or 'git' in current_category.lower():
            fit_git = "G"
        elif 'individuel' in current_category.lower() or 'fit' in current_category.lower():
            fit_git = "I"

        # Detect pricing mode from the original text
        pricing_mode, pax_count = _detect_pricing_mode(text)
        extras = section_extras.get(current_category, {"tax": Decimal("0"), "sup_single": Decimal("0")})
        tax = extras["tax"]
        sup_single = extras["sup_single"]

        # ── Calculate prices based on mode ──
        if pricing_mode == "per_pax" and (is_single or is_double or is_triple):
            # Per-person price: calculate room totals
            double_price = 2 * (price + tax)
            single_price = 1 * (price + tax) + sup_single
            twin_price = 2 * (price + tax)
        elif pricing_mode == "per_room" and pax_count >= 2:
            # Suite-style: price is for the whole room
            double_price = price if (is_double or is_suite) else None
            single_price = None
            twin_price = None
        else:
            # Per-room or unrecognized: assign to detected room type
            double_price = price if is_double or (not is_single and not is_triple) else None
            single_price = price if is_single else None
            twin_price = price if is_triple else None

        row = ParsedPriceRow(
            accommodation=description,
            city="",
            double_price=double_price,
            single_price=single_price if (is_single or is_double or is_triple) and pricing_mode == "per_pax" else single_price,
            twin_price=twin_price if (is_single or is_double or is_triple) and pricing_mode == "per_pax" else twin_price,
            hotel_type=_extract_room_type(description),
            meal_plan=meal_plan if meal_plan else None,
            fit_git=fit_git,
            season_code=None,
            note=current_category if current_category else None,
        )
        rows.append(row)

    return rows


def parse_dash_line_prices(paragraphs) -> list[ParsedPriceRow]:
    """Parse prices from dash-prefixed lines like:
        Prix agences par chambre & nuit DP/BB
        - Chambre Single :  675 DH  /595 DH
        - Chambre Double : 1155 DH /895DH

    Handles dual-price lines where two meal plans are separated by '/'.
    Also extracts hotel info, city, and contact details from surrounding text.
    """
    rows = []
    meal_plans: list[str] = []  # e.g. ["HB", "BB"] from header "DP/BB"
    extras: dict[str, dict] = {}  # extra beds etc.

    # First: detect meal plan order from header line (e.g. "DP/BB")
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # Look for header like "Prix agences par chambre & nuit DP/BB"
        header_match = re.search(
            r'(?:prix|tarif)\s+.+?\b((?:DP|BB|HB|FB|AI|RO|D\.P|B\.B|H\.B|F\.B))\s*/\s*((?:DP|BB|HB|FB|AI|RO|D\.P|B\.B|H\.B|F\.B))',
            text, re.IGNORECASE
        )
        if header_match:
            plan_map = {'DP': 'HB', 'D.P': 'HB', 'BB': 'BB', 'B.B': 'BB',
                        'HB': 'HB', 'H.B': 'HB', 'FB': 'FB', 'F.B': 'FB',
                        'AI': 'AI', 'RO': 'RO'}
            p1 = header_match.group(1).upper().replace('.', '.')
            p2 = header_match.group(2).upper().replace('.', '.')
            meal_plans = [plan_map.get(p1, p1.replace('.', '')), plan_map.get(p2, p2.replace('.', ''))]
            break

    if not meal_plans:
        # Try simpler detection: any line mentioning "DP/BB" etc.
        for p in paragraphs:
            text = p.text.strip()
            m = re.search(r'\b(DP|HB|FB)\s*/\s*(BB|RO)\b', text, re.IGNORECASE)
            if m:
                plan_map = {'DP': 'HB', 'HB': 'HB', 'FB': 'FB', 'BB': 'BB', 'RO': 'RO'}
                meal_plans = [plan_map.get(m.group(1).upper(), m.group(1).upper()),
                              plan_map.get(m.group(2).upper(), m.group(2).upper())]
                break

    # Parse price lines — group by (room_category, meal_plan), merge occupancy types
    # Key: (room_category, meal_plan) → {"single": price, "double": price, "triple": price}
    room_groups: dict[tuple[str, str | None], dict] = {}

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue

        match = DASH_PRICE_PATTERN.match(text)
        if not match:
            continue

        description = match.group(1).strip()
        price1_str = match.group(2).strip()
        price2_str = match.group(3).strip() if match.group(3) else None
        desc_lower = description.lower()

        # Skip non-room items
        if any(skip in desc_lower for skip in [
            'taxe', 't.p.t', 'tpt', 'petit', 'repas',
            'déjeuner', 'dejeuner', 'dîner', 'diner',
        ]):
            continue
        if 'lit' in desc_lower and 'sup' in desc_lower:
            continue

        price1 = parse_price(price1_str)
        price2 = parse_price(price2_str) if price2_str else None
        if not price1:
            continue

        # Detect occupancy from description
        is_single = 'single' in desc_lower
        is_double = 'double' in desc_lower
        is_triple = 'triple' in desc_lower

        # Get room category (Chambre, Suite, Suite Royale, etc.)
        room_category = _extract_room_type(description)

        # If no occupancy detected, default to double
        if not is_single and not is_double and not is_triple:
            is_double = True

        prices_by_plan = [(price1, meal_plans[0] if meal_plans else None)]
        if price2 and len(meal_plans) >= 2:
            prices_by_plan.append((price2, meal_plans[1]))
        elif price2:
            prices_by_plan.append((price2, None))

        for price, meal_plan in prices_by_plan:
            key = (room_category, meal_plan)
            if key not in room_groups:
                room_groups[key] = {}
            if is_single:
                room_groups[key]["single"] = price
            if is_double:
                room_groups[key]["double"] = price
            if is_triple:
                room_groups[key]["triple"] = price

    # Build merged rows — one per (room_category, meal_plan)
    rows = []
    for (room_category, meal_plan), prices in room_groups.items():
        rows.append(ParsedPriceRow(
            accommodation=room_category,
            city="",
            double_price=prices.get("double"),
            single_price=prices.get("single"),
            twin_price=prices.get("triple"),
            hotel_type=room_category,
            meal_plan=meal_plan,
        ))

    return rows


def _split_paragraphs_to_lines(paragraphs):
    """Split paragraph objects into individual text lines.

    Some Word documents pack multiple lines into a single paragraph
    separated by newlines. This splits them so each line is processed.
    Returns a list of fake paragraph-like objects with a .text attribute.
    """
    class FakeParagraph:
        def __init__(self, text):
            self.text = text

    lines = []
    for p in paragraphs:
        for line in p.text.split('\n'):
            lines.append(FakeParagraph(line))
    return lines


# Pattern for "1/2 chambre double en DP : 350,00 DH HT" format
HALF_ROOM_PATTERN = re.compile(
    r'^[*\-–•]?\s*1/2\s+chambre\s+(\w+)\s+en\s+(\w[\w.]*)\s*:\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b',
    re.IGNORECASE
)

# Pattern for "Supplément chambre individuelle : 250,00 DH"
SUPPLEMENT_PATTERN = re.compile(
    r'^[*\-–•]?\s*(?:suppl[ée]ment|sup)\s+(?:chambre\s+)?(?:individuelle|single)\s*:\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b',
    re.IGNORECASE
)

# Pattern for "Taxe de séjour : 16,50 DH/personne"
TAX_PATTERN = re.compile(
    r'^[*\-–•]?\s*(?:taxe\s+(?:de\s+)?s[ée]jour)\s*:\s*(\d[\d\s.,]*)\s*(?:Dhs?|DHS?|MAD)\b',
    re.IGNORECASE
)


def parse_half_room_prices(paragraphs) -> list[ParsedPriceRow]:
    """Parse "1/2 chambre double" format (per-pax pricing).

    Example:
        * Tarif GROUPE :
        1/2 chambre double en DP : 350,00 DH HT
        1/2 chambre double en BB : 260,00 DH HT
        Supplément chambre individuelle : 250,00 DH
        Taxe de séjour : 16,50 DH/personne
    """
    from decimal import Decimal

    lines = _split_paragraphs_to_lines(paragraphs)

    # Collect data per section
    sections: list[dict] = []
    current: dict | None = None

    for line in lines:
        text = line.text.strip()
        if not text:
            continue
        text_lower = text.lower()

        # Detect section headers
        section_match = re.search(
            r'(?:tarif|tarifs)\s+(.+?)(?:\s*:\s*$|\s*$)',
            text, re.IGNORECASE
        )
        if section_match:
            label = section_match.group(1).strip()
            current = {"label": label, "prices": [], "supplement": Decimal("0"), "tax": Decimal("0")}
            sections.append(current)
            continue

        if current is None:
            continue

        # Match "1/2 chambre double en DP : 350,00 DH"
        half_match = HALF_ROOM_PATTERN.match(text)
        if half_match:
            room_type = half_match.group(1).strip()  # "double"
            meal_code = half_match.group(2).strip()   # "DP" or "BB"
            price = parse_price(half_match.group(3))
            if price:
                # Map meal plan codes
                plan_map = {'DP': 'HB', 'D.P': 'HB', 'BB': 'BB', 'B.B': 'BB',
                            'HB': 'HB', 'FB': 'FB', 'AI': 'AI', 'RO': 'RO',
                            'PC': 'FB', 'P.C': 'FB'}
                meal_plan = plan_map.get(meal_code.upper(), meal_code.upper())
                current["prices"].append({"room_type": room_type.lower(), "meal_plan": meal_plan, "pax_price": price})
            continue

        # Match supplement
        sup_match = SUPPLEMENT_PATTERN.match(text)
        if sup_match:
            val = parse_price(sup_match.group(1))
            if val:
                current["supplement"] = val
            continue

        # Match tax
        tax_match = TAX_PATTERN.match(text)
        if tax_match:
            val = parse_price(tax_match.group(1))
            if val:
                current["tax"] = val
            continue

    # Build rows
    rows = []
    for section in sections:
        label = section["label"]
        supplement = section["supplement"]
        tax = section["tax"]

        fit_git = None
        label_lower = label.lower()
        if 'groupe' in label_lower or 'group' in label_lower:
            fit_git = "G"
        elif 'individuel' in label_lower or 'fit' in label_lower:
            fit_git = "I"

        season_note = None
        if 'fin' in label_lower and 'ann' in label_lower:
            season_note = "Fin d'année"

        for entry in section["prices"]:
            pax_price = entry["pax_price"]
            meal_plan = entry["meal_plan"]

            # Calculate: "1/2 chambre double" = per-pax price
            double_price = 2 * (pax_price + tax)
            single_price = 1 * (pax_price + tax) + supplement
            twin_price = 2 * (pax_price + tax)

            room_desc = f"Chambre {entry['room_type'].capitalize()}"
            row = ParsedPriceRow(
                accommodation=room_desc,
                city="",
                double_price=double_price,
                single_price=single_price,
                twin_price=twin_price,
                hotel_type=_extract_room_type(room_desc),
                meal_plan=meal_plan,
                fit_git=fit_git,
                season_code=None,
                note=f"{label}" + (f" ({season_note})" if season_note else ""),
            )
            rows.append(row)

    return rows


FRENCH_MONTHS = {
    'janvier': 1, 'jan': 1, 'janv': 1,
    'février': 2, 'fevrier': 2, 'fev': 2, 'fév': 2,
    'mars': 3, 'mar': 3,
    'avril': 4, 'avr': 4,
    'mai': 5,
    'juin': 6, 'jun': 6,
    'juillet': 7, 'juil': 7, 'jul': 7,
    'août': 8, 'aout': 8, 'aoû': 8,
    'septembre': 9, 'sep': 9, 'sept': 9,
    'octobre': 10, 'oct': 10,
    'novembre': 11, 'nov': 11,
    'décembre': 12, 'decembre': 12, 'dec': 12, 'déc': 12,
    # English
    'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5,
    'june': 6, 'july': 7, 'august': 8, 'aug': 8,
    'september': 9, 'october': 10, 'november': 11, 'december': 12,
}

# Season name normalization
SEASON_NAME_MAP = {
    'basse saison': 'Basse Saison',
    'low season': 'Basse Saison',
    'standard season': 'Basse Saison',
    'haute saison': 'Haute Saison',
    'high season': 'Haute Saison',
    'moyenne saison': 'Moyenne Saison',
    'medium season': 'Moyenne Saison',
    'mid season': 'Moyenne Saison',
    "fin d'annee": "Fin d'année",
    "fin d'année": "Fin d'année",
    'tarif normal': 'Basse Saison',
}


def _parse_date_fragment(text: str, default_year: int) -> 'date | None':
    """Parse a date from various French/English fragments.

    Handles:
        "01/11/2025", "1er Novembre 2025", "11 jan", "05 Jan",
        "01/11", "1st December 2026"
    """
    from datetime import date as date_type
    text = text.strip().rstrip('.')

    # DD/MM/YYYY
    m = re.match(r'(\d{1,2})/(\d{1,2})/(\d{4})', text)
    if m:
        try:
            return date_type(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except ValueError:
            pass

    # DD/MM (no year)
    m = re.match(r'(\d{1,2})/(\d{1,2})$', text)
    if m:
        try:
            return date_type(default_year, int(m.group(2)), int(m.group(1)))
        except ValueError:
            pass

    # DD month YYYY or DD month (French/English)
    m = re.match(
        r'(\d{1,2})(?:er|ère|e|st|nd|rd|th)?\s+([a-zéûôè]+)\.?\s*(\d{4})?',
        text, re.IGNORECASE
    )
    if m:
        day = int(m.group(1))
        month_str = m.group(2).lower().rstrip('.')
        year = int(m.group(3)) if m.group(3) else default_year
        month = FRENCH_MONTHS.get(month_str)
        if month:
            try:
                return date_type(year, month, day)
            except ValueError:
                pass

    # "month YYYY" without day (use 1st)
    m = re.match(r'([a-zéûôè]+)\.?\s*(\d{4})?$', text, re.IGNORECASE)
    if m:
        month_str = m.group(1).lower().rstrip('.')
        year = int(m.group(2)) if m.group(2) else default_year
        month = FRENCH_MONTHS.get(month_str)
        if month:
            try:
                return date_type(year, month, 1)
            except ValueError:
                pass

    return None


def _extract_season_dates(paragraphs) -> dict[str, list[ParsedDateRange]]:
    """Extract season-specific date ranges from document text.

    Looks for patterns like:
        BASSE SAISON                    HAUTE SAISON
        Du 01/11/2025 au 14/12/2025     Du 15/12/2025 au 09/01/2026
        Du 10/01/2026 au 16/02/2026     Du 19/03/2026 au 31/03/2026

    Or inline:
        BASSE SAISON • Du 10/01/2026 au 09/02/2026 • Du 10/03/2026 au 31/10/2026
        HAUTE SAISON • Du 21/12/2025 au 09/01/2026 • Du 01/03/2026 au 31/05/2026

    Or table headers:
        "Tarif Normal (MAD) (a partir de 11 jan)"
        "Tarif Période CAN (MAD) (de 22 déc au 10 jan)"

    Returns dict mapping normalized season name -> list of ParsedDateRange.
    """
    from datetime import date as date_type

    # First find default year from document
    default_year = date_type.today().year
    for p in paragraphs:
        text = p.text.strip()
        m = re.search(r'(20\d{2})\s*[/\-]\s*(20\d{2}|\d{2})\b', text)
        if m:
            y1 = int(m.group(1))
            y2s = m.group(2)
            y2 = int(y2s) if len(y2s) == 4 else int(f"20{y2s}")
            default_year = max(y1, y2)
            break

    seasons: dict[str, list[ParsedDateRange]] = {}

    # Date range pattern: "Du DD/MM/YYYY au DD/MM/YYYY" or "DD month au DD month" etc.
    date_range_re = re.compile(
        r'(?:du|from|de)\s+(.+?)\s+(?:au|to|jusqu)\s+(.+?)(?:\s*[•·,;|]|\s*$)',
        re.IGNORECASE
    )
    # "a partir de DD month" pattern
    from_re = re.compile(
        r'(?:a partir d[eu]|from)\s+(.+?)(?:\s*[•·,;|)\]]|\s*$)',
        re.IGNORECASE
    )

    # Collect all lines
    all_lines = []
    for p in paragraphs:
        text = p.text.strip()
        if text:
            for line in text.split('\n'):
                all_lines.append(line.strip())

    # Strategy 1: Look for season header followed by date lines
    current_season = None
    for line in all_lines:
        line_lower = line.lower().strip()
        if not line_lower:
            continue

        # Detect season name in line
        detected_season = None
        for key, name in SEASON_NAME_MAP.items():
            if key in line_lower:
                detected_season = name
                break

        # Check for labeled date sections like "BASSE SAISON • Du 10/01/2026 au ..."
        if detected_season:
            current_season = detected_season
            if current_season not in seasons:
                seasons[current_season] = []

        # Extract date ranges from this line
        for m in date_range_re.finditer(line):
            d_from = _parse_date_fragment(m.group(1), default_year)
            d_to = _parse_date_fragment(m.group(2), default_year)
            if d_from and d_to and current_season:
                seasons[current_season].append(
                    ParsedDateRange(date_from=d_from, date_to=d_to)
                )

        # "a partir de" pattern (open-ended: from date to Oct 31)
        if not date_range_re.search(line):
            fm = from_re.search(line)
            if fm and current_season:
                d_from = _parse_date_fragment(fm.group(1), default_year)
                if d_from:
                    d_to = date_type(default_year, 10, 31)
                    seasons[current_season].append(
                        ParsedDateRange(date_from=d_from, date_to=d_to)
                    )

    # Strategy 2: Look for "Saisonnalités" block with labeled seasons
    in_seasonality_block = False
    current_label = None
    for line in all_lines:
        line_lower = line.lower().strip()

        if 'saisonnali' in line_lower:
            in_seasonality_block = True
            continue

        if in_seasonality_block:
            # Check if this line starts a new labeled season
            for key, name in SEASON_NAME_MAP.items():
                if key in line_lower:
                    current_label = name
                    if current_label not in seasons:
                        seasons[current_label] = []
                    break

            if current_label:
                for m in date_range_re.finditer(line):
                    d_from = _parse_date_fragment(m.group(1), default_year)
                    d_to = _parse_date_fragment(m.group(2), default_year)
                    if d_from and d_to:
                        seasons[current_label].append(
                            ParsedDateRange(date_from=d_from, date_to=d_to)
                        )

                # Also handle bare "DD/MM/YYYY-DD/MM/YYYY" or "DD/MM/YYYY au DD/MM/YYYY"
                bare = re.findall(
                    r'(\d{1,2}/\d{1,2}/\d{4})\s*[-–au]+\s*(\d{1,2}/\d{1,2}/\d{4})',
                    line
                )
                for d1, d2 in bare:
                    d_from = _parse_date_fragment(d1, default_year)
                    d_to = _parse_date_fragment(d2, default_year)
                    if d_from and d_to:
                        if ParsedDateRange(date_from=d_from, date_to=d_to) not in seasons[current_label]:
                            seasons[current_label].append(
                                ParsedDateRange(date_from=d_from, date_to=d_to)
                            )

            # End block on empty or unrelated content
            if not line_lower or (not any(k in line_lower for k in ['saison', 'season', 'du ', 'from ', '/', 'ramadan', 'juillet', 'aout', 'août']) and current_label and len(seasons.get(current_label, [])) > 0):
                if not re.search(r'\d{1,2}/\d{1,2}', line):
                    in_seasonality_block = False
                    current_label = None

    return seasons


def _extract_year_range(paragraphs) -> tuple[str | None, int]:
    """Extract year range like '2025/2026' and default year from document."""
    from datetime import date as date_type
    default_year = date_type.today().year

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = re.search(r'(20\d{2})\s*[/\-]\s*(20\d{2}|\d{2})\b', text)
        if m:
            y1 = int(m.group(1))
            y2s = m.group(2)
            y2 = int(y2s) if len(y2s) == 4 else int(f"20{y2s}")
            default_year = max(y1, y2)
            return f"{y1}/{y2}", default_year

    return None, default_year


def _parse_season_table(
    data: list[list[str]],
    desc_col: int,
    season_cols: dict[int, str],
    season_table_dates: dict[str, list[ParsedDateRange]],
    fit_git: str | None = None,
) -> list[ParsedPriceRow]:
    """Parse a table with season columns like DESCRIPTION | BASSE SAISON | HAUTE SAISON.

    Handles per-pax pricing (½ Double) with supplement rows.
    Returns one row per season with calculated double/single/twin prices.
    """
    from decimal import Decimal

    # Collect per-season: base price, supplements, meal plan
    season_data: dict[str, dict] = {}
    for ci, sname in season_cols.items():
        season_data[sname] = {
            "half_double": None,
            "sup_single": Decimal("0"),
            "sup_triple": Decimal("0"),
            "sup_dp": Decimal("0"),
            "sup_pc": Decimal("0"),
            "sup_vue_mer": Decimal("0"),
            "meal_plan": None,
        }

    for row_data in data:
        if desc_col >= len(row_data):
            continue
        desc = row_data[desc_col].strip()
        desc_lower = desc.lower()
        if not desc:
            continue

        # Skip non-price rows
        if 'gratuit' in desc_lower or 'guide' in desc_lower:
            continue

        for ci, sname in season_cols.items():
            if ci >= len(row_data):
                continue
            price = parse_price(row_data[ci].strip())
            if not price:
                continue

            sd = season_data[sname]

            if '½' in desc or '1/2' in desc_lower:
                # Base per-pax price
                sd["half_double"] = price
                # Detect meal plan from description
                for plan, pattern in MEAL_PATTERNS.items():
                    if pattern.search(desc):
                        sd["meal_plan"] = plan
                        break
                if not sd["meal_plan"]:
                    if 'bb' in desc_lower or 'petit' in desc_lower:
                        sd["meal_plan"] = "BB"
                    elif 'dp' in desc_lower or 'demi' in desc_lower:
                        sd["meal_plan"] = "HB"
            elif 'single' in desc_lower or 'individuel' in desc_lower:
                if 'suppl' in desc_lower or 'sup' in desc_lower:
                    sd["sup_single"] = price
            elif 'triple' in desc_lower:
                if 'suppl' in desc_lower or 'sup' in desc_lower:
                    sd["sup_triple"] = price
            elif 'demi' in desc_lower or ('dp' in desc_lower and 'suppl' in desc_lower):
                sd["sup_dp"] = price
            elif 'pension compl' in desc_lower or ('pc' in desc_lower and 'suppl' in desc_lower):
                sd["sup_pc"] = price
            elif 'vue mer' in desc_lower or 'sea view' in desc_lower:
                sd["sup_vue_mer"] = price

    # Build rows: one per season per meal plan found
    rows: list[ParsedPriceRow] = []
    for sname, sd in season_data.items():
        if sd["half_double"] is None:
            continue
        pax_price = sd["half_double"]
        # double = 2 × pax_price, single = 2 × pax_price + sup_single, twin = 2 × pax_price + sup_triple
        double_price = 2 * pax_price
        single_price = 2 * pax_price + sd["sup_single"] if sd["sup_single"] else double_price
        twin_price = 2 * pax_price + sd["sup_triple"] if sd["sup_triple"] else None

        date_ranges = season_table_dates.get(sname, [])

        rows.append(ParsedPriceRow(
            accommodation="Chambre",
            city="",
            double_price=double_price,
            single_price=single_price,
            twin_price=twin_price,
            hotel_type="Chambre",
            meal_plan=sd["meal_plan"],
            fit_git=fit_git,
            date_ranges=list(date_ranges),
            note=sname,
        ))

        # If DP supplement exists, also create an HB row
        if sd["sup_dp"] and sd["meal_plan"] == "BB":
            dp_pax = pax_price + sd["sup_dp"]
            rows.append(ParsedPriceRow(
                accommodation="Chambre",
                city="",
                double_price=2 * dp_pax,
                single_price=2 * dp_pax + sd["sup_single"] if sd["sup_single"] else 2 * dp_pax,
                twin_price=2 * dp_pax + sd["sup_triple"] if sd["sup_triple"] else None,
                hotel_type="Chambre",
                meal_plan="HB",
                fit_git=fit_git,
                date_ranges=list(date_ranges),
                note=sname,
            ))

        # If PC supplement exists, also create an FB row
        if sd["sup_pc"] and sd["meal_plan"] == "BB":
            pc_pax = pax_price + sd["sup_pc"]
            rows.append(ParsedPriceRow(
                accommodation="Chambre",
                city="",
                double_price=2 * pc_pax,
                single_price=2 * pc_pax + sd["sup_single"] if sd["sup_single"] else 2 * pc_pax,
                twin_price=2 * pc_pax + sd["sup_triple"] if sd["sup_triple"] else None,
                hotel_type="Chambre",
                meal_plan="FB",
                fit_git=fit_git,
                date_ranges=list(date_ranges),
                note=sname,
            ))

    return rows


class WordParser(BaseParser):
    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in (".docx", ".doc")

    def parse(self, file_path: Path) -> list[ParsedPriceRow]:
        all_rows = []
        doc = DocxDocument(file_path)

        # Split paragraphs into individual lines (some docs pack multiple lines in one paragraph)
        lines = _split_paragraphs_to_lines(doc.paragraphs)

        # Also collect text from all table cells for season date extraction
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        for line in text.split('\n'):
                            if line.strip():
                                table_lines.append(type('P', (), {'text': line.strip()})())
        # Combine paragraph lines + table lines for season scanning
        all_text_lines = lines + table_lines

        # Also extract season dates from dedicated season tables
        # (tables where rows map season name → date ranges)
        season_table_dates: dict[str, list[ParsedDateRange]] = {}
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    cell0_lower = cells[0].lower()
                    # Check if first cell is a season name
                    season_name = None
                    for key, name in SEASON_NAME_MAP.items():
                        if key in cell0_lower:
                            season_name = name
                            break
                    if season_name:
                        # Parse date ranges from remaining cells
                        combined = " ".join(cells[1:])
                        from datetime import date as _dt
                        default_year = _dt.today().year
                        # Find default year from document
                        for p in lines:
                            m = re.search(r'(20\d{2})\s*[/\-]\s*(20\d{2}|\d{2})\b', p.text)
                            if m:
                                y2s = m.group(2)
                                default_year = int(y2s) if len(y2s) == 4 else int(f"20{y2s}")
                                break
                        # Extract "Du X au Y" patterns
                        for m in re.finditer(
                            r'[Dd]u\s+(\d{1,2}/\d{1,2}/\d{4})\s+au\s*(\d{1,2}/\d{1,2}/\d{4})',
                            combined
                        ):
                            d_from = _parse_date_fragment(m.group(1), default_year)
                            d_to = _parse_date_fragment(m.group(2), default_year)
                            if d_from and d_to:
                                if season_name not in season_table_dates:
                                    season_table_dates[season_name] = []
                                season_table_dates[season_name].append(
                                    ParsedDateRange(date_from=d_from, date_to=d_to)
                                )

        # First try: parse tables in document order, tracking FIT/GIT context
        # Build ordered list of (type, element) from document body
        from docx.oxml.ns import qn
        current_fit_git = None
        table_index = 0
        table_map = {id(t._tbl): t for t in doc.tables}

        for element in doc.element.body:
            # Track FIT/GIT context from paragraphs
            if element.tag == qn('w:p'):
                text = element.text or ""
                # Also get text from runs
                for run in element.findall(qn('w:r')):
                    t = run.find(qn('w:t'))
                    if t is not None and t.text:
                        text += t.text
                text_lower = text.strip().lower()
                if text_lower:
                    if any(k in text_lower for k in ['tour operat', 'tarifs individu', 'tarif individu', 'tarif fit', 'tarifs fit']):
                        current_fit_git = "I"
                    elif any(k in text_lower for k in ['tarifs groupe', 'tarif groupe', 'tarifs git', 'tarif git', 'groupe applicable']):
                        current_fit_git = "G"
                    elif any(k in text_lower for k in ['saisonnali', 'réduction', 'reduction', 'conditions', 'article']):
                        # Reset context when we hit non-pricing sections
                        pass

            # Process tables
            if element.tag == qn('w:tbl'):
                tbl_obj = table_map.get(id(element))
                if not tbl_obj or len(tbl_obj.rows) < 2:
                    continue

                has_content = False
                for row in tbl_obj.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            has_content = True
                            break
                    if has_content:
                        break
                if not has_content:
                    continue

                header = [cell.text.strip() for cell in tbl_obj.rows[0].cells]
                header_lower = [h.lower() for h in header]
                data = []
                for row in tbl_obj.rows[1:]:
                    data.append([cell.text.strip() for cell in row.cells])

                # Check if this is a season-column table (DESCRIPTION | BASSE SAISON | HAUTE SAISON)
                season_cols = {}
                for ci, hl in enumerate(header_lower):
                    for key, name in SEASON_NAME_MAP.items():
                        if key in hl:
                            season_cols[ci] = name
                            break
                desc_col = None
                for ci, hl in enumerate(header_lower):
                    if 'description' in hl or 'type' in hl or 'chambre' in hl or 'room' in hl:
                        desc_col = ci
                        break
                if desc_col is None and len(header) > 0 and 0 not in season_cols:
                    desc_col = 0

                if season_cols and desc_col is not None and len(header) >= 2:
                    season_rows = _parse_season_table(data, desc_col, season_cols, season_table_dates, fit_git=current_fit_git)
                    all_rows.extend(season_rows)
                elif len(header) >= 3:
                    df = pd.DataFrame(data, columns=header)
                    temp_xlsx = file_path.with_suffix(".tmp.xlsx")
                    df.to_excel(temp_xlsx, index=False)
                    try:
                        parser = ExcelParser()
                        rows = parser.parse(temp_xlsx)
                        # Apply current FIT/GIT context to parsed rows
                        for r in rows:
                            if current_fit_git and not r.fit_git:
                                r.fit_git = current_fit_git
                        all_rows.extend(rows)
                    finally:
                        temp_xlsx.unlink(missing_ok=True)

        # Second try: dotted-line format ("Chambre Double ……… 600 Dhs / TTC")
        if not all_rows:
            rows = parse_unstructured_prices(lines)
            all_rows.extend(rows)

        # Third try: dash-line format ("- Chambre Single : 675 DH /595 DH")
        if not all_rows:
            rows = parse_dash_line_prices(lines)
            all_rows.extend(rows)

        # Fourth try: half-room format ("1/2 chambre double en DP : 350,00 DH")
        if not all_rows:
            rows = parse_half_room_prices(lines)
            all_rows.extend(rows)

        # Enrich rows with hotel info from document headers/body
        if all_rows:
            hotel_name, city, hotel_stars, address, phone, email_addr = extract_hotel_info_from_doc(doc)
            # Also try to find hotel name in body text if not in headers
            if hotel_name == "Unknown Hotel":
                hotel_name, city, body_stars, phone = _extract_hotel_from_body(doc.paragraphs, city)
                if body_stars and not hotel_stars:
                    hotel_stars = body_stars

            # Extract season date ranges (from paragraphs + table text)
            season_dates = _extract_season_dates(all_text_lines)
            year_range, default_year = _extract_year_range(all_text_lines)
            # Merge in season dates found in dedicated season tables
            for sname, sranges in season_table_dates.items():
                if sname not in season_dates:
                    season_dates[sname] = sranges
                else:
                    # Add any ranges not already present
                    existing = {(str(r.date_from), str(r.date_to)) for r in season_dates[sname]}
                    for r in sranges:
                        if (str(r.date_from), str(r.date_to)) not in existing:
                            season_dates[sname].append(r)

            for row in all_rows:
                if city and not row.city:
                    row.city = city
                if hotel_stars and not row.stars:
                    row.stars = hotel_stars
                if address and not row.address:
                    row.address = address
                if phone and not row.phone:
                    row.phone = phone
                if email_addr and not row.email:
                    row.email = email_addr
                if hotel_name and hotel_name != "Unknown Hotel":
                    # Set accommodation to hotel name; room type belongs in hotel_type
                    if not row.accommodation.startswith(hotel_name):
                        row.accommodation = hotel_name

                # Apply season info
                if year_range and not row.season_code:
                    row.season_code = year_range

                if not row.date_ranges:
                    if season_dates:
                        # Try to match row's note/season_code to a season
                        row_note = (row.note or "").lower()
                        row_season = (row.season_code or "").lower()
                        matched = False
                        for season_name, ranges in season_dates.items():
                            sn_lower = season_name.lower()
                            if sn_lower in row_note or sn_lower in row_season:
                                row.date_ranges = list(ranges)
                                matched = True
                                break

                        if not matched:
                            # If only one season found, apply to all rows
                            if len(season_dates) == 1:
                                only_ranges = list(season_dates.values())[0]
                                row.date_ranges = list(only_ranges)
                            # If seasons found but row doesn't match, duplicate row per season
                            elif len(season_dates) > 1:
                                # Don't assign — will be handled by row duplication below
                                pass
                    elif year_range:
                        # Fallback: today to Oct 31 of newest year
                        from datetime import date as date_type
                        row.date_ranges = [
                            ParsedDateRange(
                                date_from=date_type.today(),
                                date_to=date_type(default_year, 10, 31),
                            )
                        ]

            # If multiple seasons found, duplicate rows that have no date_ranges
            # (one copy per season, with season name in note)
            if season_dates and len(season_dates) > 1:
                expanded_rows = []
                for row in all_rows:
                    if row.date_ranges:
                        expanded_rows.append(row)
                    else:
                        for season_name, ranges in season_dates.items():
                            import copy
                            new_row = copy.copy(row)
                            new_row.date_ranges = list(ranges)
                            existing_note = new_row.note or ""
                            if season_name.lower() not in existing_note.lower():
                                new_row.note = f"{existing_note} ({season_name})".strip().lstrip("(").rstrip(")") if not existing_note else f"{existing_note} ({season_name})"
                                if not existing_note:
                                    new_row.note = season_name
                            expanded_rows.append(new_row)
                all_rows = expanded_rows

        return all_rows
