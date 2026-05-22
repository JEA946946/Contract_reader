import hashlib
import json
import shutil
from concurrent.futures import ThreadPoolExecutor
from datetime import date
from decimal import Decimal
from typing import List, Optional
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.database import get_db
from app.config import settings
from app.models import Document
from app.schemas import (
    DocumentOut,
    UploadForReviewResponse,
    ParsedRowSchema,
    ParsedDateRangeSchema,
    ParsedMenuRowSchema,
    ParsedTransportRowSchema,
    ConfirmRowsRequest,
    ConfirmRowsResponse,
    ConfirmMenuRowsRequest,
    ConfirmTransportRowsRequest,
)
from app.utils import detect_file_type
from app.services.extraction import (
    parse_document, parse_document_background, parse_restaurant_document, rows_to_json,
    deduplicate_rows, deduplicate_menu_rows, _menu_row_to_dict, _transport_row_to_dict,
)
from app.services.normalizer import save_parsed_rows
from app.services.restaurant_normalizer import save_parsed_menu_rows
from app.services.transport_normalizer import save_parsed_transport_rows
from app.services.folder_service import auto_assign_folder
from app.parsers.base import ParsedPriceRow, ParsedMenuRow, ParsedDateRange, ParsedTransportRow

router = APIRouter(prefix="/api/documents", tags=["documents"])

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "xlsx", "xls", "csv", "txt", "eml", "msg"}

# Module-level worker pool for background parsing (Phase 3.1)
_parse_executor = ThreadPoolExecutor(max_workers=3, thread_name_prefix="doc-parse")


def _compute_content_hash(file_path: Path) -> str:
    """Compute SHA-256 hash of file contents for duplicate detection."""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _extract_entity_name(doc: Document) -> Optional[str]:
    """Extract unique hotel/restaurant name(s) from cached parsed rows, including city."""
    if not doc.parsed_rows_json:
        return None
    try:
        rows = json.loads(doc.parsed_rows_json)
        if doc.document_category == "restaurant":
            names = list(dict.fromkeys(r.get("restaurant_name") for r in rows if r.get("restaurant_name")))
        elif doc.document_category == "transportation":
            names = list(dict.fromkeys(r.get("company_name") for r in rows if r.get("company_name")))
        else:
            names = list(dict.fromkeys(r.get("accommodation") for r in rows if r.get("accommodation")))
        if not names:
            return None
        label = ", ".join(names)
        # Append city from first row that has one
        for row in rows:
            city = (row.get("city") or "").strip()
            if city:
                label = f"{label}, {city}"
                break
        return label
    except Exception:
        return None


def _doc_to_out(doc: Document) -> dict:
    """Convert a Document to a DocumentOut-compatible dict with hotel_name."""
    return DocumentOut(
        id=doc.id,
        filename=doc.filename,
        file_type=doc.file_type,
        upload_date=doc.upload_date,
        status=doc.status,
        row_count=doc.row_count,
        notes=doc.notes,
        hotel_name=_extract_entity_name(doc),
        document_category=doc.document_category or "hotel",
        folder_id=doc.folder_id,
        confidence_score=float(doc.confidence_score) if doc.confidence_score is not None else None,
    )


def _mark_orphaned_documents(docs: list, db: Session) -> None:
    """Mark documents whose source file no longer exists on disk as failed."""
    changed = False
    for doc in docs:
        upload_path = settings.upload_path / doc.filename
        if not upload_path.exists():
            doc.status = "failed"
            doc.notes = "Source file no longer exists on disk"
            changed = True
    if changed:
        db.commit()


def _row_to_schema(row: ParsedPriceRow) -> ParsedRowSchema:
    """Convert a ParsedPriceRow dataclass to a Pydantic schema for JSON."""
    return ParsedRowSchema(
        accommodation=row.accommodation or "",
        city=row.city or "",
        room_desc=row.room_desc,
        double_price=float(row.double_price) if row.double_price is not None else None,
        single_price=float(row.single_price) if row.single_price is not None else None,
        twin_price=float(row.twin_price) if row.twin_price is not None else None,
        triple_price=float(row.triple_price) if row.triple_price is not None else None,
        quadruple_price=float(row.quadruple_price) if row.quadruple_price is not None else None,
        stars=row.stars,
        hotel_type=row.hotel_type,
        meal_plan=row.meal_plan,
        fit_git=row.fit_git,
        season_code=row.season_code,
        baby_discount=row.baby_discount,
        child_discount=row.child_discount,
        date_ranges=[
            ParsedDateRangeSchema(
                date_from=dr.date_from.isoformat() if dr.date_from else None,
                date_to=dr.date_to.isoformat() if dr.date_to else None,
            )
            for dr in row.date_ranges
        ],
        min_stay=row.min_stay,
        note=row.note,
        address=row.address,
        phone=row.phone,
        email=row.email,
    )


def _menu_row_to_schema(row: ParsedMenuRow) -> ParsedMenuRowSchema:
    """Convert a ParsedMenuRow dataclass to a Pydantic schema."""
    return ParsedMenuRowSchema(
        restaurant_name=row.restaurant_name or "",
        city=row.city or "",
        menu_name=row.menu_name,
        description=row.description,
        lunch_price=float(row.lunch_price) if row.lunch_price is not None else None,
        dinner_price=float(row.dinner_price) if row.dinner_price is not None else None,
        lunch_child_price=float(row.lunch_child_price) if row.lunch_child_price is not None else None,
        dinner_child_price=float(row.dinner_child_price) if row.dinner_child_price is not None else None,
        course_1=row.course_1,
        course_2=row.course_2,
        course_3=row.course_3,
        course_4=row.course_4,
        course_5=row.course_5,
        min_pax=row.min_pax,
        drink_included=row.drink_included,
        season_code=row.season_code,
        date_ranges=[
            ParsedDateRangeSchema(
                date_from=dr.date_from.isoformat() if dr.date_from else None,
                date_to=dr.date_to.isoformat() if dr.date_to else None,
            )
            for dr in row.date_ranges
        ],
        note=row.note,
        address=row.address,
        phone=row.phone,
        email=row.email,
    )


def _schema_to_menu_row(s: ParsedMenuRowSchema) -> ParsedMenuRow:
    """Convert a Pydantic schema back to a ParsedMenuRow dataclass."""
    return ParsedMenuRow(
        restaurant_name=s.restaurant_name or "",
        city=s.city or "",
        menu_name=s.menu_name,
        description=s.description,
        lunch_price=Decimal(str(s.lunch_price)) if s.lunch_price is not None else None,
        dinner_price=Decimal(str(s.dinner_price)) if s.dinner_price is not None else None,
        lunch_child_price=Decimal(str(s.lunch_child_price)) if s.lunch_child_price is not None else None,
        dinner_child_price=Decimal(str(s.dinner_child_price)) if s.dinner_child_price is not None else None,
        course_1=s.course_1,
        course_2=s.course_2,
        course_3=s.course_3,
        course_4=s.course_4,
        course_5=s.course_5,
        min_pax=s.min_pax,
        drink_included=s.drink_included,
        season_code=s.season_code,
        date_ranges=[
            ParsedDateRange(
                date_from=date.fromisoformat(dr.date_from) if dr.date_from else None,
                date_to=date.fromisoformat(dr.date_to) if dr.date_to else None,
            )
            for dr in s.date_ranges
        ],
        note=s.note,
        address=s.address,
        phone=s.phone,
        email=s.email,
    )


def _schema_to_row(s: ParsedRowSchema) -> ParsedPriceRow:
    """Convert a Pydantic schema back to a ParsedPriceRow dataclass."""
    return ParsedPriceRow(
        accommodation=s.accommodation,
        city=s.city,
        room_desc=s.room_desc,
        double_price=Decimal(str(s.double_price)) if s.double_price is not None else None,
        single_price=Decimal(str(s.single_price)) if s.single_price is not None else None,
        twin_price=Decimal(str(s.twin_price)) if s.twin_price is not None else None,
        triple_price=Decimal(str(s.triple_price)) if s.triple_price is not None else None,
        quadruple_price=Decimal(str(s.quadruple_price)) if s.quadruple_price is not None else None,
        stars=s.stars,
        hotel_type=s.hotel_type,
        meal_plan=s.meal_plan,
        fit_git=s.fit_git,
        season_code=s.season_code,
        baby_discount=s.baby_discount,
        child_discount=s.child_discount,
        date_ranges=[
            ParsedDateRange(
                date_from=date.fromisoformat(dr.date_from) if dr.date_from else None,
                date_to=date.fromisoformat(dr.date_to) if dr.date_to else None,
            )
            for dr in s.date_ranges
        ],
        min_stay=s.min_stay,
        note=s.note,
        address=s.address,
        phone=s.phone,
        email=s.email,
    )


def _transport_row_to_schema(row: ParsedTransportRow) -> ParsedTransportRowSchema:
    """Convert a ParsedTransportRow dataclass to a Pydantic schema."""
    return ParsedTransportRowSchema(
        code=row.code or "",
        price=float(row.price) if row.price is not None else None,
        company_name=row.company_name or "",
        company_code=row.company_code or "",
        product=row.product or "",
        bus_size=row.bus_size,
        service_type=row.service_type,
        days=row.days,
        route_description=row.route_description or "",
        note=row.note or "",
        city=row.city or "",
    )


def _schema_to_transport_row(s: ParsedTransportRowSchema) -> ParsedTransportRow:
    """Convert a Pydantic schema back to a ParsedTransportRow dataclass."""
    return ParsedTransportRow(
        code=s.code or "",
        price=Decimal(str(s.price)) if s.price is not None else None,
        company_name=s.company_name or "",
        company_code=s.company_code or "",
        product=s.product or "",
        bus_size=s.bus_size,
        service_type=s.service_type or "",
        days=s.days,
        route_description=s.route_description or "",
        note=s.note or "",
        city=s.city or "",
    )


@router.post("/upload", response_model=UploadForReviewResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    file_type = detect_file_type(file.filename)
    if file_type not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # Save file
    upload_path = settings.upload_path / file.filename
    with open(upload_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Compute content hash for duplicate detection (Phase 1.3)
    content_hash = _compute_content_hash(upload_path)

    # Check for existing document with the same content hash
    existing = (
        db.query(Document)
        .filter(Document.content_hash == content_hash)
        .filter(Document.status.in_(("pending_review", "completed")))
        .first()
    )
    if existing and existing.parsed_rows_json:
        # Duplicate detected — clone the result instead of re-parsing
        document = Document(
            filename=file.filename,
            file_type=file_type,
            status=existing.status,
            content_hash=content_hash,
            row_count=existing.row_count,
            parsed_rows_json=existing.parsed_rows_json,
            document_category=existing.document_category,
            confidence_score=existing.confidence_score,
            notes=f"Duplicate of document #{existing.id} — cached result reused",
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        return UploadForReviewResponse(
            id=document.id,
            filename=document.filename,
            status=document.status,
            row_count=document.row_count or 0,
            message=f"Duplicate detected (matches doc #{existing.id}) — cached result reused",
            rows=[],
            menu_rows=[],
            document_category=document.document_category or "hotel",
        )

    # Create document record
    document = Document(
        filename=file.filename,
        file_type=file_type,
        status="processing",
        content_hash=content_hash,
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    # Submit to worker pool (Phase 3.1) — replaces ad-hoc threading.Thread
    _parse_executor.submit(parse_document_background, document.id)

    return UploadForReviewResponse(
        id=document.id,
        filename=document.filename,
        status="processing",
        row_count=0,
        message="Document uploaded — processing in background",
        rows=[],
        menu_rows=[],
        document_category="hotel",
    )


@router.post("/{document_id}/confirm", response_model=ConfirmRowsResponse)
def confirm_document(document_id: int, body: ConfirmRowsRequest, db: Session = Depends(get_db)):
    from app.services.feedback_analyzer import save_feedback

    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if document.status not in ("pending_review", "processing"):
        raise HTTPException(status_code=400, detail=f"Document status is '{document.status}', expected 'pending_review'")

    # Save feedback diff if rows were modified (Phase 2.2)
    original_json = document.parsed_rows_json
    parsed_rows = [_schema_to_row(s) for s in body.rows]
    corrected_json = rows_to_json(parsed_rows)
    if original_json:
        save_feedback(document.id, original_json, corrected_json, db)

    # Update cached JSON so auto-folder uses the user-edited data
    document.parsed_rows_json = corrected_json
    count = save_parsed_rows(parsed_rows, document, db)
    document.row_count = count
    document.status = "completed"
    auto_assign_folder(document, db)
    db.commit()

    return ConfirmRowsResponse(
        id=document.id,
        filename=document.filename,
        status=document.status,
        row_count=count,
        message=f"Saved {count} price rows",
    )


@router.post("/{document_id}/confirm-restaurant", response_model=ConfirmRowsResponse)
def confirm_restaurant_document(document_id: int, body: ConfirmMenuRowsRequest, db: Session = Depends(get_db)):
    from app.services.feedback_analyzer import save_feedback

    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if document.status not in ("pending_review", "processing"):
        raise HTTPException(status_code=400, detail=f"Document status is '{document.status}', expected 'pending_review'")

    original_json = document.parsed_rows_json
    parsed_rows = [_schema_to_menu_row(s) for s in body.rows]
    corrected_json = json.dumps([_menu_row_to_dict(r) for r in parsed_rows])
    if original_json:
        save_feedback(document.id, original_json, corrected_json, db)

    document.parsed_rows_json = corrected_json
    count = save_parsed_menu_rows(parsed_rows, document, db)
    document.row_count = count
    document.status = "completed"
    document.document_category = "restaurant"
    auto_assign_folder(document, db)
    db.commit()

    return ConfirmRowsResponse(
        id=document.id,
        filename=document.filename,
        status=document.status,
        row_count=count,
        message=f"Saved {count} menu price rows",
    )


@router.post("/{document_id}/confirm-transportation", response_model=ConfirmRowsResponse)
def confirm_transportation_document(document_id: int, body: ConfirmTransportRowsRequest, db: Session = Depends(get_db)):
    from app.services.feedback_analyzer import save_feedback

    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if document.status not in ("pending_review", "processing"):
        raise HTTPException(status_code=400, detail=f"Document status is '{document.status}', expected 'pending_review'")

    original_json = document.parsed_rows_json
    parsed_rows = [_schema_to_transport_row(s) for s in body.rows]
    corrected_json = json.dumps([_transport_row_to_dict(r) for r in parsed_rows])
    if original_json:
        save_feedback(document.id, original_json, corrected_json, db)

    document.parsed_rows_json = corrected_json
    count = save_parsed_transport_rows(parsed_rows, document, db)
    document.row_count = count
    document.status = "completed"
    document.document_category = "transportation"
    auto_assign_folder(document, db)
    db.commit()

    return ConfirmRowsResponse(
        id=document.id,
        filename=document.filename,
        status=document.status,
        row_count=count,
        message=f"Saved {count} transport price rows",
    )


@router.get("/pending", response_model=List[DocumentOut])
def list_pending_documents(category: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(Document).filter(Document.status == "pending_review")
    if category:
        query = query.filter(Document.document_category == category)
    docs = query.order_by(Document.upload_date.desc()).all()
    _mark_orphaned_documents(docs, db)
    return [_doc_to_out(d) for d in docs if d.status == "pending_review"]


@router.get("/{document_id}/parsed-rows", response_model=UploadForReviewResponse)
def get_parsed_rows(document_id: int, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if document.status not in ("pending_review", "processing", "completed"):
        raise HTTPException(
            status_code=400,
            detail=f"Document status is '{document.status}', not reviewable",
        )

    category = document.document_category or "hotel"

    # If still processing, return empty rows with processing status
    if document.status == "processing":
        return UploadForReviewResponse(
            id=document.id,
            filename=document.filename,
            status="processing",
            row_count=0,
            message="Document is still being processed",
            rows=[],
            menu_rows=[],
            transport_rows=[],
            document_category=category,
        )

    # Serve from cache if available
    if document.parsed_rows_json:
        if category == "restaurant":
            raw_rows = json.loads(document.parsed_rows_json)
            # Migrate old-format rows (menu_tag/price_per_person) to new format
            if raw_rows and ("menu_tag" in raw_rows[0] or "price_per_person" in raw_rows[0]):
                for r in raw_rows:
                    if "price_per_person" in r and r["price_per_person"] is not None:
                        tag = (r.get("menu_tag") or "").lower()
                        if "lunch" in tag:
                            r.setdefault("lunch_price", r["price_per_person"])
                        else:
                            r.setdefault("dinner_price", r["price_per_person"])
                    for old_key in ("menu_tag", "price_per_person"):
                        r.pop(old_key, None)
                    for new_key in ("lunch_price", "dinner_price", "lunch_child_price", "dinner_child_price",
                                    "course_1", "course_2", "course_3", "course_4", "course_5"):
                        r.setdefault(new_key, None)
                # Update the cached JSON with migrated data
                document.parsed_rows_json = json.dumps(raw_rows)
                db.commit()
            cached_menu_rows = [ParsedMenuRowSchema(**r) for r in raw_rows]
            return UploadForReviewResponse(
                id=document.id,
                filename=document.filename,
                status=document.status,
                row_count=len(cached_menu_rows),
                message=f"{len(cached_menu_rows)} menu rows ready for review",
                rows=[],
                menu_rows=cached_menu_rows,
                document_category=category,
            )
        elif category == "transportation":
            cached_transport_rows = [ParsedTransportRowSchema(**r) for r in json.loads(document.parsed_rows_json)]
            return UploadForReviewResponse(
                id=document.id,
                filename=document.filename,
                status=document.status,
                row_count=len(cached_transport_rows),
                message=f"{len(cached_transport_rows)} transport rows ready for review",
                rows=[],
                transport_rows=cached_transport_rows,
                document_category=category,
            )
        else:
            cached_rows = [ParsedRowSchema(**r) for r in json.loads(document.parsed_rows_json)]
            return UploadForReviewResponse(
                id=document.id,
                filename=document.filename,
                status=document.status,
                row_count=len(cached_rows),
                message=f"{len(cached_rows)} price rows ready for review",
                rows=cached_rows,
                menu_rows=[],
                document_category=category,
            )

    # Fallback for legacy docs without cache: re-parse and cache
    upload_path = settings.upload_path / document.filename
    if not upload_path.exists():
        raise HTTPException(status_code=404, detail="Source file not found on disk")

    rows = parse_document(upload_path, document, db)
    # parse_document() may re-detect category and cache results in parsed_rows_json
    category = document.document_category or "hotel"
    if category not in ("restaurant", "transportation"):
        document.parsed_rows_json = rows_to_json(rows)
    db.commit()

    # For restaurant/transport, read back from the cache that parse_document() just set
    if category == "restaurant" and document.parsed_rows_json:
        cached_menu_rows = [ParsedMenuRowSchema(**r) for r in json.loads(document.parsed_rows_json)]
        return UploadForReviewResponse(
            id=document.id,
            filename=document.filename,
            status=document.status,
            row_count=len(cached_menu_rows),
            message=f"Re-parsed {len(cached_menu_rows)} menu rows",
            rows=[],
            menu_rows=cached_menu_rows,
            document_category=category,
        )

    if category == "transportation" and document.parsed_rows_json:
        cached_transport_rows = [ParsedTransportRowSchema(**r) for r in json.loads(document.parsed_rows_json)]
        return UploadForReviewResponse(
            id=document.id,
            filename=document.filename,
            status=document.status,
            row_count=len(cached_transport_rows),
            message=f"Re-parsed {len(cached_transport_rows)} transport rows",
            rows=[],
            transport_rows=cached_transport_rows,
            document_category=category,
        )

    return UploadForReviewResponse(
        id=document.id,
        filename=document.filename,
        status=document.status,
        row_count=len(rows),
        message=f"Re-parsed {len(rows)} price rows",
        rows=[_row_to_schema(r) for r in rows],
        menu_rows=[],
        document_category=category,
    )


@router.get("/{document_id}/file")
def get_document_file(document_id: int, db: Session = Depends(get_db)):
    """Serve the original uploaded document file."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    upload_path = settings.upload_path / document.filename
    if not upload_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    import mimetypes
    content_type = mimetypes.guess_type(document.filename)[0] or "application/octet-stream"

    return FileResponse(
        path=upload_path,
        media_type=content_type,
        headers={"Content-Disposition": f'inline; filename="{document.filename}"'},
    )


@router.get("/{document_id}/attachments")
def list_document_attachments(document_id: int, db: Session = Depends(get_db)):
    """List attachments in an email document (.eml/.msg)."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    upload_path = settings.upload_path / document.filename
    if not upload_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    if document.file_type not in ("eml", "msg"):
        # Non-email: return the document itself as the single "attachment"
        return [{"index": 0, "filename": document.filename, "content_type": "application/pdf" if document.file_type == "pdf" else "application/octet-stream"}]

    import email
    from email import policy as email_policy
    msg = email.message_from_bytes(upload_path.read_bytes(), policy=email_policy.default)
    attachments = []
    idx = 0
    for part in msg.walk():
        disposition = str(part.get("Content-Disposition", ""))
        if "attachment" in disposition or (part.get_content_type() == "application/pdf" and part.get_filename()):
            fname = part.get_filename() or f"attachment_{idx}"
            import mimetypes
            ct = part.get_content_type() or mimetypes.guess_type(fname)[0] or "application/octet-stream"
            attachments.append({"index": idx, "filename": fname, "content_type": ct})
            idx += 1
    return attachments


@router.get("/{document_id}/attachments/{att_index}")
def get_document_attachment(document_id: int, att_index: int, db: Session = Depends(get_db)):
    """Serve a specific attachment from an email document."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    upload_path = settings.upload_path / document.filename
    if not upload_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    if document.file_type not in ("eml", "msg"):
        # Non-email: serve the file itself if index 0
        if att_index != 0:
            raise HTTPException(status_code=404, detail="Attachment not found")
        import mimetypes
        content_type = mimetypes.guess_type(document.filename)[0] or "application/octet-stream"
        return FileResponse(path=upload_path, media_type=content_type, headers={"Content-Disposition": f'inline; filename="{document.filename}"'})

    import email
    from email import policy as email_policy
    msg = email.message_from_bytes(upload_path.read_bytes(), policy=email_policy.default)
    idx = 0
    for part in msg.walk():
        disposition = str(part.get("Content-Disposition", ""))
        if "attachment" in disposition or (part.get_content_type() == "application/pdf" and part.get_filename()):
            if idx == att_index:
                fname = part.get_filename() or f"attachment_{idx}"
                data = part.get_payload(decode=True)
                if not data:
                    raise HTTPException(status_code=404, detail="Empty attachment")
                import mimetypes
                import tempfile
                ct = part.get_content_type() or mimetypes.guess_type(fname)[0] or "application/octet-stream"
                suffix = Path(fname).suffix
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                tmp.write(data)
                tmp.close()
                return FileResponse(path=tmp.name, media_type=ct, headers={"Content-Disposition": f'inline; filename="{fname}"'})
            idx += 1
    raise HTTPException(status_code=404, detail="Attachment not found")


@router.post("/{document_id}/ai-reparse", response_model=UploadForReviewResponse)
def ai_reparse_document(document_id: int, db: Session = Depends(get_db)):
    """Force re-parse a document using the category-appropriate AI parser."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    upload_path = settings.upload_path / document.filename
    if not upload_path.exists():
        raise HTTPException(status_code=404, detail="Source file not found on disk")

    file_type = detect_file_type(document.filename)

    # Re-detect category on reparse to fix misclassified documents
    from app.services.extraction import _detect_category
    detected = _detect_category(upload_path, file_type)
    document.document_category = detected
    category = detected

    try:
        if category == "restaurant":
            menu_rows = parse_restaurant_document(upload_path, document, db)
            menu_rows = deduplicate_menu_rows(menu_rows)
            document.status = "pending_review"
            document.row_count = len(menu_rows)
            document.notes = f"AI re-parsed (restaurant): {len(menu_rows)} rows"
            document.parsed_rows_json = rows_to_json(menu_rows)
            db.commit()
            return UploadForReviewResponse(
                id=document.id,
                filename=document.filename,
                status=document.status,
                row_count=len(menu_rows),
                message=f"AI extracted {len(menu_rows)} menu price rows",
                rows=[],
                menu_rows=[_menu_row_to_schema(r) for r in menu_rows],
                document_category="restaurant",
            )

        if category == "transportation":
            from app.services.extraction import parse_transport_document
            transport_rows = parse_transport_document(upload_path, document, db)
            document.status = "pending_review"
            document.row_count = len(transport_rows)
            document.notes = f"AI re-parsed (transport): {len(transport_rows)} rows"
            document.parsed_rows_json = rows_to_json(transport_rows)
            db.commit()
            return UploadForReviewResponse(
                id=document.id,
                filename=document.filename,
                status=document.status,
                row_count=len(transport_rows),
                message=f"AI extracted {len(transport_rows)} transport rows",
                rows=[],
                transport_rows=[_transport_row_to_schema(r) for r in transport_rows],
                document_category="transportation",
            )

        # Default: hotel
        from app.parsers.ai_parser import AiParser
        ai_parser = AiParser()

        if file_type in ("pdf",):
            rows = ai_parser.parse_pdf(upload_path)
        elif file_type in ("eml", "msg"):
            from app.parsers.email_parser import EmailParser
            rows = EmailParser().parse(upload_path)
        elif file_type in ("docx", "doc"):
            from docx import Document as DocxDocument
            doc = DocxDocument(upload_path)
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        text_parts.append(" | ".join(cells))
            rows = ai_parser.parse_text("\n".join(text_parts))
        elif file_type in ("txt", "csv", "text"):
            rows = ai_parser.parse_text(upload_path.read_text(encoding="utf-8", errors="replace"))
        elif file_type in ("xlsx", "xls"):
            import pandas as pd
            dfs = pd.read_excel(upload_path, sheet_name=None, header=None)
            text_parts = []
            for sheet_name, df in dfs.items():
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                for _, row in df.iterrows():
                    cells = [str(c) if pd.notna(c) else "" for c in row]
                    if any(cells):
                        text_parts.append(" | ".join(cells))
            rows = ai_parser.parse_text("\n".join(text_parts))
        else:
            raise HTTPException(status_code=400, detail=f"AI reparse not supported for '{file_type}' files")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI parsing failed: {str(e)}")

    from app.services.extraction import collapse_hotel_rows
    rows = deduplicate_rows(rows)
    rows = collapse_hotel_rows(rows)
    document.status = "pending_review"
    document.row_count = len(rows)
    document.notes = f"AI re-parsed: {len(rows)} rows"
    document.parsed_rows_json = rows_to_json(rows)
    db.commit()

    return UploadForReviewResponse(
        id=document.id,
        filename=document.filename,
        status=document.status,
        row_count=len(rows),
        message=f"AI extracted {len(rows)} price rows",
        rows=[_row_to_schema(r) for r in rows],
        document_category="hotel",
    )


class ParseTextRequest(BaseModel):
    text: str


@router.post("/parse-text", response_model=UploadForReviewResponse)
def parse_text_input(body: ParseTextRequest, db: Session = Depends(get_db)):
    """Parse pasted text through the AI parser and create a document for review."""
    from app.parsers.ai_parser import AiParser

    text = body.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="No text provided")

    # Create document record for the pasted text
    document = Document(
        filename="pasted-text.txt",
        file_type="txt",
        status="processing",
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    # Give unique filename and save text to disk so review/re-parse works
    document.filename = f"pasted-text-{document.id}.txt"
    db.commit()
    upload_path = settings.upload_path / document.filename
    upload_path.write_text(text, encoding="utf-8")

    try:
        ai_parser = AiParser()
        rows = ai_parser.parse_text(text)
        rows = deduplicate_rows(rows)

        document.status = "pending_review"
        document.row_count = len(rows)
        document.notes = f"Parsed from pasted text: {len(rows)} rows"
        document.parsed_rows_json = rows_to_json(rows)
        db.commit()

        return UploadForReviewResponse(
            id=document.id,
            filename=document.filename,
            status=document.status,
            row_count=len(rows),
            message=f"Parsed {len(rows)} price rows from text — ready for review",
            rows=[_row_to_schema(r) for r in rows],
        )
    except Exception as e:
        document.status = "failed"
        document.notes = str(e)
        db.commit()
        return UploadForReviewResponse(
            id=document.id,
            filename=document.filename,
            status="failed",
            row_count=0,
            message=f"Processing failed: {str(e)}",
            rows=[],
        )


@router.get("", response_model=List[DocumentOut])
def list_documents(category: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(Document)
    if category:
        query = query.filter(Document.document_category == category)
    docs = query.order_by(Document.upload_date.desc()).all()
    _mark_orphaned_documents(
        [d for d in docs if d.status in ("pending_review", "processing")], db
    )
    return [_doc_to_out(d) for d in docs]


@router.delete("/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db)):
    from app.models import Hotel, Restaurant, TransportCompany, ProcessedEmail

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Null out all FK references so constraints don't block delete
    db.query(Hotel).filter(Hotel.source_document_id == document_id).update(
        {"source_document_id": None}, synchronize_session="fetch"
    )
    db.query(Restaurant).filter(Restaurant.source_document_id == document_id).update(
        {"source_document_id": None}, synchronize_session="fetch"
    )
    db.query(TransportCompany).filter(TransportCompany.source_document_id == document_id).update(
        {"source_document_id": None}, synchronize_session="fetch"
    )
    db.query(ProcessedEmail).filter(ProcessedEmail.document_id == document_id).update(
        {"document_id": None}, synchronize_session="fetch"
    )

    # Delete file
    upload_path = settings.upload_path / doc.filename
    upload_path.unlink(missing_ok=True)

    db.delete(doc)
    db.commit()
    return {"message": "Document deleted"}


class BatchReparseRequest(BaseModel):
    filter: str  # "failed" or "zero_rows"


@router.post("/batch-reparse")
def batch_reparse_documents(body: BatchReparseRequest, db: Session = Depends(get_db)):
    """Batch reprocess failed or zero-row documents."""
    from app.services.auto_retry import MAX_RETRIES, _get_retry_count

    BATCH_CAP = 20

    if body.filter == "failed":
        candidates = (
            db.query(Document)
            .filter(Document.status == "failed")
            .order_by(Document.upload_date.desc())
            .limit(BATCH_CAP * 2)
            .all()
        )
        # Filter by retry count
        eligible = []
        for doc in candidates:
            if _get_retry_count(doc.notes) < MAX_RETRIES:
                upload_path = settings.upload_path / doc.filename
                if upload_path.exists():
                    eligible.append(doc)
            if len(eligible) >= BATCH_CAP:
                break
        skipped = len(candidates) - len(eligible)

    elif body.filter == "zero_rows":
        candidates = (
            db.query(Document)
            .filter(
                Document.status.in_(("pending_review", "completed", "failed")),
                (Document.row_count == 0) | (Document.row_count.is_(None)),
            )
            .order_by(Document.upload_date.desc())
            .limit(BATCH_CAP * 2)
            .all()
        )
        eligible = []
        for doc in candidates:
            upload_path = settings.upload_path / doc.filename
            if upload_path.exists():
                eligible.append(doc)
            if len(eligible) >= BATCH_CAP:
                break
        skipped = len(candidates) - len(eligible)

    else:
        raise HTTPException(status_code=400, detail="filter must be 'failed' or 'zero_rows'")

    # Queue each for reprocessing
    for doc in eligible:
        doc.status = "processing"
        doc.notes = (doc.notes or "") + " | Batch reprocess"
    db.commit()

    for doc in eligible:
        _parse_executor.submit(parse_document_background, doc.id)

    return {
        "queued": len(eligible),
        "skipped": skipped,
        "message": f"Queued {len(eligible)} documents for reprocessing, skipped {skipped}",
    }


class UpdateDocumentEntityRequest(BaseModel):
    entity_name: str
    city: str


@router.patch("/{document_id}/entity", response_model=DocumentOut)
def update_document_entity(document_id: int, body: UpdateDocumentEntityRequest, db: Session = Depends(get_db)):
    """Update the entity name and city in a document's parsed rows, and re-assign folder."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    category = doc.document_category or "hotel"

    if doc.parsed_rows_json:
        rows = json.loads(doc.parsed_rows_json)
    else:
        # Create a minimal row so the entity name and city are stored
        rows = [{}]

    # Update name and city in all rows
    for row in rows:
        if category == "restaurant":
            row["restaurant_name"] = body.entity_name
        elif category == "transportation":
            row["company_name"] = body.entity_name
        else:
            row["accommodation"] = body.entity_name
        row["city"] = body.city

    doc.parsed_rows_json = json.dumps(rows)
    # Clear folder so auto-assign can re-evaluate based on the new city
    if body.city.strip():
        doc.folder_id = None
    auto_assign_folder(doc, db)
    db.commit()

    return _doc_to_out(doc)
